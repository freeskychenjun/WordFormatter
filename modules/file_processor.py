import os
import re
import shutil
import tempfile
import win32com.client
from docx import Document

from .exception_handler import FileProcessingError, global_exception_handler


class FileProcessor:
    def __init__(self, log_callback=None):
        self.temp_files = []
        self.log_callback = log_callback
        self.com_app = None

    def _log(self, message):
        if self.log_callback:
            self.log_callback(message)

    def _cleanup_temp_files(self):
        self._log("正在清理本轮临时文件...")
        for f in self.temp_files:
            try:
                if os.path.exists(f):
                    os.remove(f)
                    self._log(f"  > 临时文件 {os.path.basename(f)} 已删除")
            except OSError as e:
                self._log(f"  > 警告：删除临时文件 {f} 失败: {e}")
        self.temp_files.clear()

    def _get_wps_app(self):
        if self.com_app is None:
            self._log("首次需要，正在启动WPS/Word应用...")
            try:
                self.com_app = win32com.client.Dispatch('KWPS.Application')
                self._log("  > 已成功连接到WPS。")
            except Exception:
                try:
                    self.com_app = win32com.client.Dispatch('Word.Application')
                    self._log("  > 已成功连接到Word。")
                except Exception as e:
                    raise RuntimeError(f"未能启动WPS或Word，请确保已安装。错误: {e}")
            self.com_app.Visible = False
        return self.com_app

    def quit_com_app(self):
        if self.com_app:
            self._log("所有任务完成，正在关闭WPS/Word应用...")
            self.com_app.Quit()
            self.com_app = None
            self._log("  > 应用已关闭。")

    def convert_to_docx(self, input_path):
        try:
            file_ext = os.path.splitext(input_path)[1].lower()
            is_from_txt = (file_ext == '.txt')

            # 使用系统临时目录来避免权限问题
            temp_dir = tempfile.gettempdir()
            base_name = os.path.splitext(os.path.basename(input_path))[0]

            # 清理文件名中的特殊字符，避免在Windows系统中出现问题
            # 移除或替换可能引起问题的字符
            cleaned_base_name = re.sub(r'[<>:"/\\|?*\x00-\x1F]', '_', base_name)
            # 限制文件名长度，避免超过系统限制
            cleaned_base_name = cleaned_base_name[:100] if len(cleaned_base_name) > 100 else cleaned_base_name
            # 移除末尾的点和空格（Windows不允许）
            cleaned_base_name = cleaned_base_name.rstrip('. ')

            # 添加调试信息
            self._log(f"  > 原始文件路径: {input_path}")
            self._log(f"  > 系统临时目录: {temp_dir}")
            self._log(f"  > 清理后的文件名: {cleaned_base_name}")

            if file_ext == '.docx':
                self._log("检测到 .docx 文件，正在创建安全的处理副本...")
                temp_docx_path = os.path.join(temp_dir, f"~temp_copy_{cleaned_base_name}.docx")
                self._log(f"  > 临时文件路径: {temp_docx_path}")

                # 检查目标路径是否存在同名文件，如果存在则删除
                if os.path.exists(temp_docx_path):
                    self._log(f"  > 检测到同名临时文件，正在删除: {temp_docx_path}")
                    try:
                        os.remove(temp_docx_path)
                    except Exception as e:
                        self._log(f"  > 删除同名临时文件失败: {e}")

                try:
                    shutil.copy2(input_path, temp_docx_path)
                    self.temp_files.append(temp_docx_path)
                    self._log(f"  > 副本创建成功: {os.path.basename(temp_docx_path)}")

                    # 验证文件是否创建成功
                    if os.path.exists(temp_docx_path):
                        self._log(f"  > 验证文件存在: {temp_docx_path}")
                    else:
                        self._log(f"  > 警告：文件创建后不存在: {temp_docx_path}")

                    return temp_docx_path, is_from_txt
                except Exception as e:
                    self._log(f"  > 创建副本失败: {e}")
                    # 如果复制失败，尝试使用不同的临时文件名
                    for i in range(10):
                        alt_temp_docx_path = os.path.join(temp_dir, f"~temp_copy_{cleaned_base_name}_{i}.docx")
                        self._log(f"  > 尝试备用路径: {alt_temp_docx_path}")
                        if not os.path.exists(alt_temp_docx_path):
                            try:
                                shutil.copy2(input_path, alt_temp_docx_path)
                                self.temp_files.append(alt_temp_docx_path)
                                self._log(f"  > 使用备用名称创建副本成功: {os.path.basename(alt_temp_docx_path)}")
                                return alt_temp_docx_path, is_from_txt
                            except Exception as e2:
                                self._log(f"  > 备用名称创建副本也失败: {e2}")
                                continue
                    # 如果所有尝试都失败，抛出异常
                    raise FileProcessingError(f"无法创建文件副本: {e}", e)

            temp_docx_path = os.path.join(temp_dir, f"~temp_converted_{cleaned_base_name}.docx")
            self.temp_files.append(temp_docx_path)
            self._log(f"  > 转换文件路径: {temp_docx_path}")

            if file_ext == '.txt':
                self._log("检测到 .txt 文件，正在创建 .docx...")
                doc = Document()
                try:
                    with open(input_path, 'r', encoding='utf-8') as f:
                        for line in f:
                            doc.add_paragraph(line.strip())
                    self._log("  > 已使用 UTF-8 编码读取TXT文件。")
                except UnicodeDecodeError:
                    self._log("  > UTF-8读取失败，尝试使用 GBK 编码...")
                    with open(input_path, 'r', encoding='gbk') as f:
                        for line in f:
                            doc.add_paragraph(line.strip())
                    self._log("  > 已成功使用 GBK 编码读取TXT文件。")
                doc.save(temp_docx_path)
                self._log("TXT转换完成。")

                # 验证文件是否创建成功
                if os.path.exists(temp_docx_path):
                    self._log(f"  > 验证文件存在: {temp_docx_path}")
                else:
                    self._log(f"  > 警告：文件创建后不存在: {temp_docx_path}")

                return temp_docx_path, is_from_txt
            elif file_ext in ['.wps', '.doc']:
                self._log(f"正在转换 {file_ext} 文件为 .docx...")
                app = self._get_wps_app()
                doc_com = app.Documents.Open(os.path.abspath(input_path), ReadOnly=1)
                doc_com.SaveAs2(os.path.abspath(temp_docx_path), FileFormat=12)
                doc_com.Close()
                self._log("文件格式转换完成。")

                # 验证文件是否创建成功
                if os.path.exists(temp_docx_path):
                    self._log(f"  > 验证文件存在: {temp_docx_path}")
                else:
                    self._log(f"  > 警告：文件创建后不存在: {temp_docx_path}")

                return temp_docx_path, is_from_txt

            raise FileProcessingError(f"不支持的文件格式: {file_ext}")
        except Exception as e:
            error_msg = global_exception_handler.handle_exception(e, "文件转换")
            self._log(f"文件转换过程中发生错误: {error_msg}")
            raise

    def _preprocess_com_tasks(self, docx_path):
        self._log("正在对副本执行预处理（接受所有修订、转换自动编号）...")
        app = self._get_wps_app()
        try:
            doc_com = app.Documents.Open(os.path.abspath(docx_path))

            doc_com.TrackRevisions = False
            self._log("  > 已关闭修订追踪。")

            if doc_com.Revisions.Count > 0:
                doc_com.AcceptAllRevisions()
                self._log("  > 已接受文档副本中的所有修订。")

            doc_com.Content.ListFormat.ConvertNumbersToText()
            self._log("  > 已将副本中的自动编号转换为文本。")

            if doc_com.Revisions.Count > 0:
                doc_com.AcceptAllRevisions()
                self._log("  > 已接受编号转换产生的修订。")

            doc_com.TrackRevisions = False

            # 确保文档被正确保存
            doc_com.Save()
            doc_com.Close()
            self._log("预处理完成。")

            # 添加延迟以确保文件操作完成
            import time
            time.sleep(0.5)

            # 验证文件是否存在
            if os.path.exists(docx_path):
                self._log(f"  > 验证预处理后文件存在: {docx_path}")
            else:
                self._log(f"  > 警告：预处理后文件不存在: {docx_path}")
        except Exception as e:
            self._log(f"警告：执行预处理任务时出错: {e}")
