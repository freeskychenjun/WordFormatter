from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import RGBColor


class PageSetup:
    def __init__(self, config, log_callback=None):
        self.config = config
        self.log_callback = log_callback

    def _log(self, message):
        if self.log_callback:
            self.log_callback(message)

    def _set_run_font(self, run, font_name, size_pt, set_color=False, is_bold=False):
        """设置单个run的字体属性"""
        # 尝试多种字体名称设置方式
        # 1. 设置高级API的font.name属性
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.bold = is_bold

        if set_color:
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # 2. 直接设置XML的rFonts元素，确保所有字符类型都使用相同字体
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        
        # 设置所有字符类型的字体
        rFonts.set(qn('w:ascii'), font_name)      # ASCII字符（英文、数字、符号）
        rFonts.set(qn('w:hAnsi'), font_name)      # 高ASCII字符
        rFonts.set(qn('w:eastAsia'), font_name)   # 中文字体
        
        # 3. 额外设置cs（复杂脚本）字体，确保所有语言都能正确显示
        rFonts.set(qn('w:cs'), font_name)
        
        # 4. 清除可能存在的字体主题设置，确保使用指定的字体
        if hasattr(rFonts, 'themeFont'):
            rFonts.themeFont = None
        if hasattr(rFonts, 'themeFontAscii'):
            rFonts.themeFontAscii = None
        if hasattr(rFonts, 'themeFontHAnsi'):
            rFonts.themeFontHAnsi = None
        if hasattr(rFonts, 'themeFontEastAsia'):
            rFonts.themeFontEastAsia = None
        if hasattr(rFonts, 'themeFontCs'):
            rFonts.themeFontCs = None

    def _create_page_number(self, paragraph, text):
        font_name = self.config['page_number_font']
        font_size = self.config['page_number_size']
        self._set_run_font(paragraph.add_run('— '), font_name, font_size, set_color=True)
        run_field = paragraph.add_run()
        self._set_run_font(run_field, font_name, font_size, set_color=True)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = text
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_field._r.extend([fldChar1, instrText, fldChar2])
        self._set_run_font(paragraph.add_run(' —'), font_name, font_size, set_color=True)

    def _apply_page_setup(self, doc, is_from_txt=False):
        self._log("正在应用页面边距设置...")

        # 判断是否需要强制设置A4纸
        # 逻辑：如果是纯文本来源（包括直接输入）则设置为A4
        should_set_a4 = is_from_txt

        for section in doc.sections:
            section.top_margin = Cm(self.config['margin_top'])
            section.bottom_margin = Cm(self.config['margin_bottom'])
            section.left_margin = Cm(self.config['margin_left'])
            section.right_margin = Cm(self.config['margin_right'])

            # 设置纸张大小为A4 (仅在需要时)
            if should_set_a4:
                section.page_width = Cm(21)
                section.page_height = Cm(29.7)

        if should_set_a4:
            self._log("  > 已将页面大小设置为 A4。")