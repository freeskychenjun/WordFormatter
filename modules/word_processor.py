import logging
import os
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

from .file_processor import FileProcessor
from .document_formatter import DocumentFormatter
from .title_handler import TitleHandler
from .page_setup import PageSetup
from .config_manager import ConfigManager

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


class WordProcessor:
    def __init__(self, config, log_callback=None):
        self.config = config
        self.log_callback = log_callback
        self.file_processor = FileProcessor(log_callback)
        self.document_formatter = DocumentFormatter(config, log_callback)
        self.title_handler = TitleHandler(config, log_callback)
        self.page_setup = PageSetup(config, log_callback)

    def _log(self, message):
        if self.log_callback:
            self.log_callback(message)

    def format_document(self, input_path, output_path):
        processing_path, is_from_txt = self.file_processor.convert_to_docx(input_path)
        self._log(f"  > 处理路径: {processing_path}")

        # 检查临时文件是否存在
        if not os.path.exists(processing_path):
            self._log(f"  > 错误：临时文件不存在: {processing_path}")
            raise FileNotFoundError(f"临时文件不存在: {processing_path}")

        if not is_from_txt:
            self.file_processor._preprocess_com_tasks(processing_path)
            # 预处理完成后，再次检查文件是否存在
            if not os.path.exists(processing_path):
                self._log(f"  > 警告：预处理后文件不存在: {processing_path}")
                # 尝试重新创建文件
                self._log(f"  > 尝试重新创建文件...")
                processing_path, _ = self.file_processor.convert_to_docx(input_path)
                if not os.path.exists(processing_path):
                    self._log(f"  > 重新创建文件也失败: {processing_path}")
                    raise FileNotFoundError(f"预处理后文件不存在且重新创建失败: {processing_path}")
                else:
                    self._log(f"  > 重新创建文件成功: {processing_path}")
            else:
                self._log(f"  > 验证预处理后文件仍存在: {processing_path}")

        # 再次检查文件是否存在
        if not os.path.exists(processing_path):
            self._log(f"  > 错误：预处理后文件不存在: {processing_path}")
            raise FileNotFoundError(f"预处理后文件不存在: {processing_path}")

        # 尝试打开文档，如果失败则尝试重新创建
        doc = None
        try:
            doc = Document(processing_path)
        except Exception as e:
            self._log(f"  > 首次打开文档失败: {e}")
            # 尝试使用绝对路径
            abs_path = os.path.abspath(processing_path)
            self._log(f"  > 尝试使用绝对路径: {abs_path}")
            try:
                doc = Document(abs_path)
            except Exception as e2:
                self._log(f"  > 使用绝对路径也失败: {e2}")
                # 如果还是失败，尝试重新创建文件
                self._log(f"  > 尝试重新创建文件...")
                processing_path, _ = self.file_processor.convert_to_docx(input_path)
                if os.path.exists(processing_path):
                    self._log(f"  > 重新创建文件成功，再次尝试打开...")
                    try:
                        doc = Document(processing_path)
                        self._log(f"  > 重新打开文档成功")
                    except Exception as e3:
                        self._log(f"  > 重新打开文档也失败: {e3}")
                        raise
                else:
                    self._log(f"  > 重新创建文件失败: {processing_path}")
                    raise FileNotFoundError(f"重新创建文件失败: {processing_path}")

        if doc is None:
            raise RuntimeError("无法打开文档")

        all_blocks = list(self.document_formatter._iter_block_items(doc))
        processed_indices = set()

        apply_color = not is_from_txt

        if not is_from_txt:
            self._log("正在扫描图表标题...")
            for idx, block in enumerate(all_blocks):
                is_pic_para = isinstance(block, Paragraph) and (
                            '<w:drawing>' in block._p.xml or '<w:pict>' in block._p.xml)
                is_table = isinstance(block, Table)

                if not (is_pic_para or is_table): continue

                for direction in [-1, 1]:
                    caption_found = False
                    for i in range(idx + direction, -1 if direction == -1 else len(all_blocks), direction):
                        if i in processed_indices: continue
                        potential_caption = all_blocks[i]
                        if not isinstance(potential_caption, Paragraph): break
                        text = potential_caption.text.strip()
                        if text:
                            # 单独成行且以"图"或"表"开头的段落都识别为图表标题，不再要求居中对齐
                            # 特别地，对于表格标题，即使不是居中对齐也会被识别
                            if text.startswith("图") or text.startswith("表"):
                                detected_type = "图" if text.startswith("图") else "表"
                                # 如果标题未居中，设置为居中对齐
                                if potential_caption.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                                    potential_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    self._log(f"  > 已将未居中的{detected_type}标题设置为居中对齐")
                                self._log(f"  > 发现 {detected_type} 的标题: \"{text[:30]}...\" (在段落 {i + 1})")
                                config_font_key = f'{("figure" if detected_type == "图" else "table")}_caption_font'
                                config_size_key = f'{("figure" if detected_type == "图" else "table")}_caption_size'
                                config_font = self.config[config_font_key]
                                config_size = self.config[config_size_key]
                                config_bold_key = f'{("figure" if detected_type == "图" else "table")}_caption_bold'
                                config_bold = self.config.get(config_bold_key, False)
                                self.document_formatter._apply_font_to_runs(potential_caption, config_font, config_size,
                                                                           set_color=apply_color, is_bold=config_bold)
                                # 表格/图表标题不缩进，确保完全没有任何缩进
                                potential_caption.paragraph_format.first_line_indent = None
                                potential_caption.paragraph_format.left_indent = Pt(0)
                                potential_caption.paragraph_format.right_indent = Pt(0)
                                # 移除 hanging_indent 设置，因为该属性不存在

                                # 确保完全清除任何可能的缩进设置 - 使用更健壮的方式
                                try:
                                    pPr = potential_caption._p.get_or_add_pPr()

                                    # 获取或创建缩进元素
                                    if pPr.find(qn('w:ind')) is None:
                                        ind = OxmlElement('w:ind')
                                        pPr.append(ind)
                                    else:
                                        ind = pPr.find(qn('w:ind'))

                                    # 清除所有可能的缩进属性
                                    for attr in ['w:firstLine', 'w:firstLineChars', 'w:left', 'w:leftChars',
                                                 'w:right', 'w:rightChars', 'w:hanging', 'w:hangingChars']:
                                        if attr in ind.attrib:
                                            del ind.attrib[attr]

                                    # 显式设置为0 - 使用不同单位确保彻底移除缩进
                                    ind.set(qn('w:firstLineChars'), '0')
                                    ind.set(qn('w:leftChars'), '0')
                                    ind.set(qn('w:rightChars'), '0')
                                    ind.set(qn('w:firstLine'), '0')
                                    ind.set(qn('w:left'), '0')
                                    ind.set(qn('w:right'), '0')
                                except Exception as e:
                                    self._log(f"  > 设置 {detected_type} 标题缩进时出错: {e}")

                                # 额外保障措施：记录当前处理的标题，防止后续被覆盖
                                # 使用 setattr 来动态添加属性
                                setattr(potential_caption, '_has_no_indent', True)

                                # 应用大纲级别设置
                                outline_level_key = f'{("figure" if detected_type == "图" else "table")}_caption_outline_level'
                                if outline_level_key in self.config:
                                    outline_level_value = self.config[outline_level_key]
                                    if outline_level_value != '无' and outline_level_value != '':
                                        try:
                                            level = int(outline_level_value)
                                            if 1 <= level <= 9:
                                                self.document_formatter._set_outline_level(potential_caption, level)
                                                self._log(f"  > 已设置 {detected_type} 标题的大纲级别为 {level}")
                                        except (ValueError, TypeError):
                                            pass  # 忽略无效值

                                processed_indices.add(i)
                                caption_found = True
                            break
                    if caption_found: break

        # 不再查找主标题和副标题
        title_indices, subtitle_indices = [], []

        self._log("预扫描完成，开始逐段格式化...")
        if self.config['set_outline']:
            self._log("【大纲级别设置已启用】")
        else:
            self._log("【大纲级别设置已禁用】")

        re_h1 = re.compile(r'^[一二三四五六七八九十百千万零]+\s*、')
        re_h2 = re.compile(r'^[（\(][一二三四五六七八九十百千万零]+[）\)]')
        # 原有的正则表达式保留，但添加更精确的数字编号识别
        re_h3 = re.compile(r'^\d+\s*[\.．]')
        re_h4 = re.compile(r'^[（\(]\d+[）\)]')
        # 添加新的正则表达式，根据点的数量识别不同级别的数字编号标题
        re_number_h2 = re.compile(r'^\d+\s*[\.．]\s*\d+\s*[\.．]\s*$')  # 例如 "7.9." (1个点)
        re_number_h3 = re.compile(r'^\d+\s*[\.．]\s*\d+\s*[\.．]\s*\d+\s*[\.．]\s*$')  # 例如 "7.9.4." (2个点)
        re_number_h4 = re.compile(r'^\d+\s*[\.．]\s*\d+\s*[\.．]\s*\d+\s*[\.．]\s*\d+\s*[\.．]\s*$')  # 例如 "7.8.5.3." (3个点)
        re_number_h5 = re.compile(r'^\d+\s*[\.．]\s*\d+\s*[\.．]\s*\d+\s*[\.．]\s*\d+\s*[\.．]\s*\d+\s*[\.．]\s*$')  # 例如 "7.8.5.3.1." (4个点)
        re_attachment = re.compile(r'^附件\s*(\d+|[一二三四五六七八九十百千万零]+)?\s*[:：]?$')

        # 已移除主标题和副标题的格式化功能
        # 格式化主标题
        # if title_indices:
        #     self._log(f"\n开始格式化主标题（共 {len(title_indices)} 行）...")
        #     for idx in title_indices:
        #         para = all_blocks[idx]
        #         self._log(f"段落 {idx + 1}: 主标题行 - \"{para.text[:30]}...\"")
        #         self.document_formatter._strip_leading_whitespace(para)
        #         self.document_formatter._apply_font_to_runs(para, self.config['h1_font'], self.config['h1_size'],
        #                                                     set_color=apply_color)
        #         para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #         para.paragraph_format.first_line_indent = None
        #         para.paragraph_format.left_indent = Pt(0)
        #
        #         # 设置标题行间距
        #         spacing = para._p.get_or_add_pPr().get_or_add_spacing()
        #         spacing.set(qn('w:beforeAutospacing'), '0')
        #         spacing.set(qn('w:afterAutospacing'), '0')
        #         # 根据配置设置一级标题段前、段后间距（直接使用磅值）
        #         h1_space_before_pts = float(self.config['h1_space_before'])
        #         h1_space_after_pts = float(self.config['h1_space_after'])
        #         para.paragraph_format.space_before = Pt(h1_space_before_pts)
        #         para.paragraph_format.space_after = Pt(h1_space_after_pts)
        #         para.paragraph_format.line_spacing = Pt(self.config['line_spacing'])
        #         
        #         # 设置大纲级别为1级
        #         if self.config['set_outline']:
        #             self.document_formatter._set_outline_level(para, 1)
        #             self._log(f"  > 已设置主标题的大纲级别为 1")
        #
        #         self.document_formatter._reset_pagination_properties(para)
        #
        # # 格式化副标题
        # if subtitle_indices:
        #     self._log(f"\n开始格式化副标题（共 {len(subtitle_indices)} 行）...")
        #     for idx in subtitle_indices:
        #         para = all_blocks[idx]
        #         self._log(f"段落 {idx + 1}: 副标题行 - \"{para.text[:30]}...\"")
        #         self.document_formatter._strip_leading_whitespace(para)
        #         self.document_formatter._apply_font_to_runs(para, self.config['h2_font'], self.config['h2_size'],
        #                                                     set_color=apply_color)
        #         para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #         para.paragraph_format.first_line_indent = None
        #         para.paragraph_format.left_indent = Pt(0)
        #
        #         # 设置副标题行间距
        #         spacing = para._p.get_or_add_pPr().get_or_add_spacing()
        #         spacing.set(qn('w:beforeAutospacing'), '0')
        #         spacing.set(qn('w:afterAutospacing'), '0')
        #         # 根据配置设置二级标题段前、段后间距（直接使用磅值）
        #         h2_space_before_pts = float(self.config['h2_space_before'])
        #         h2_space_after_pts = float(self.config['h2_space_after'])
        #         para.paragraph_format.space_before = Pt(h2_space_before_pts)
        #         para.paragraph_format.space_after = Pt(h2_space_after_pts)
        #         para.paragraph_format.line_spacing = Pt(self.config['line_spacing'])
        #
        #         self.document_formatter._reset_pagination_properties(para)

        block_idx = 0
        while block_idx < len(all_blocks):
            block = all_blocks[block_idx]

            if block_idx in processed_indices:
                self._log(f"块 {block_idx + 1}: 已作为图表/附件标题处理 - 跳过")
                block_idx += 1
                continue

            current_block_num = block_idx + 1
            if isinstance(block, Table):
                self._log(f"块 {current_block_num}: 表格 - 检查内部标题")
                # 检查表格内部第一行是否为标题
                table = block
                if len(table.rows) > 0:
                    first_row = table.rows[0]
                    # 检查第一行的所有单元格
                    for cell in first_row.cells:
                        if len(cell.paragraphs) > 0:
                            for para in cell.paragraphs:
                                text = para.text.strip()
                                if text and text.startswith("表"):
                                    # 确保标题始终居中对齐
                                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    self._log(f"  > 发现表格内部标题: \"{text[:30]}...\"")
                                    config_font = self.config['table_caption_font']
                                    config_size = self.config['table_caption_size']
                                    config_bold = self.config.get('table_caption_bold', False)
                                    self.document_formatter._apply_font_to_runs(para, config_font, config_size,
                                                                        set_color=apply_color,
                                                                        is_bold=config_bold)
                                    # 表格标题不缩进，确保完全没有任何缩进
                                    para.paragraph_format.first_line_indent = None
                                    para.paragraph_format.left_indent = Pt(0)
                                    para.paragraph_format.right_indent = Pt(0)
                                    para.paragraph_format.hanging_indent = Pt(0)
                                    
                                    # 确保彻底清除所有缩进相关设置 - 使用更健壮的方式
                                    try:
                                        pPr = para._p.get_or_add_pPr()
                                        
                                        # 完全移除现有的缩进元素，然后创建一个全新的
                                        existing_ind = pPr.find(qn('w:ind'))
                                        if existing_ind is not None:
                                            pPr.remove(existing_ind)
                                        
                                        # 创建一个全新的、干净的缩进元素
                                        new_ind = OxmlElement('w:ind')
                                        
                                        # 显式设置所有可能的缩进属性为0，使用多种单位确保彻底移除缩进
                                        new_ind.set(qn('w:firstLineChars'), '0')
                                        new_ind.set(qn('w:leftChars'), '0')
                                        new_ind.set(qn('w:rightChars'), '0')
                                        new_ind.set(qn('w:firstLine'), '0')
                                        new_ind.set(qn('w:left'), '0')
                                        new_ind.set(qn('w:right'), '0')
                                        new_ind.set(qn('w:hanging'), '0')
                                        new_ind.set(qn('w:hangingChars'), '0')
                                        
                                        # 添加到pPr元素
                                        pPr.append(new_ind)
                                        
                                        # 检查并移除可能影响缩进的其他元素
                                        for element_name in ['w:firstLine', 'w:firstLineChars', 'w:left', 'w:leftChars',
                                                            'w:right', 'w:rightChars', 'w:hanging', 'w:hangingChars']:
                                            element = pPr.find(qn(element_name))
                                            if element is not None:
                                                pPr.remove(element)
                                        
                                        # 设置标记，表明这个段落已经明确设置为无缩进
                                        para._has_no_indent = True
                                        
                                    except Exception as e:
                                        self._log(f"设置表格标题缩进时出错: {e}")
                                        # 即使发生异常，仍然尝试通过简单的API调用确保没有缩进
                                        try:
                                            para.paragraph_format.first_line_indent = None
                                            para.paragraph_format.left_indent = Pt(0)
                                            # 额外添加这一行，直接设置为负数以强制覆盖可能存在的缩进
                                            para.paragraph_format.left_indent = Pt(-0.01)
                                            # 然后再设置为0
                                            para.paragraph_format.left_indent = Pt(0)
                                        except:
                                            pass
                                    
                                    # 标记为表格标题，避免后续被覆盖
                                    para._is_table_caption = True

                                # 应用表格标题大纲级别设置
                                if 'table_caption_outline_level' in self.config:
                                    outline_level_value = self.config['table_caption_outline_level']
                                    if outline_level_value != '无' and outline_level_value != '':
                                        try:
                                            level = int(outline_level_value)
                                            if 1 <= level <= 9:
                                                self.document_formatter._set_outline_level(para, level)
                                                self._log(f"  > 已设置表格内部标题的大纲级别为 {level}")
                                        except (ValueError, TypeError):
                                            pass  # 忽略无效值

                                break
                # 处理表格内的所有单元格内容
                body_font = self.config['body_font']
                
                # 遍历表格的所有行和单元格，确保表格内容的字体大小始终保持不变
                for row in table.rows:
                    for cell in row.cells:
                        # 遍历单元格中的所有段落
                        for para in cell.paragraphs:
                            # 跳过已经处理过的表格标题
                            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER and para.text.strip().startswith("表"):
                                continue
                                 
                            # 检查是否为表格标题
                            is_table_caption = hasattr(para, '_is_table_caption') and para._is_table_caption
                            if is_table_caption:
                                continue
                             
                            # 遍历段落中的所有run
                            for run in para.runs:
                                # 保存原始字体大小
                                original_font_size = run.font.size
                                
                                # 确保表格使用的字体与正文一致，同时保持字体大小不变
                                if original_font_size:
                                    # 仅修改字体名称为正文字体，严格使用原始字体大小
                                    self.document_formatter._set_run_font(
                                        run, body_font, 
                                        original_font_size.pt,  # 严格使用原始字体大小
                                        set_color=apply_color,
                                        use_times_roman_for_ascii=self.config.get('table_use_times_roman', True)
                                    )
                                elif not original_font_size:
                                    # 对于没有明确字体大小的情况，使用正文大小作为默认值
                                    self.document_formatter._set_run_font(
                                        run, body_font, 
                                        self.config['body_size'],  # 使用正文字体大小作为默认值
                                        set_color=apply_color,
                                        use_times_roman_for_ascii=self.config.get('table_use_times_roman', True)
                                    )
                                # 其他情况不做任何修改，完全保持表格内容的原始格式
                block_idx += 1
                continue

            para = block
            if not para.text.strip():
                self._log(f"段落 {current_block_num}: 空白 - 跳过")
                block_idx += 1
                continue

            is_pic = '<w:drawing>' in para._p.xml or '<w:pict>' in para._p.xml
            is_embedded_obj = '<w:object>' in para._p.xml
            if is_pic or is_embedded_obj:
                log_msg = "图片" if is_pic else "附件"
                self._log(f"段落 {current_block_num}: {log_msg} - 仅格式化文字")

                text_to_check = para.text.lstrip()
                para_text_preview = text_to_check[:30].replace("\n", " ")

                if re_h1.match(text_to_check):
                    self._log(f"  > 文字识别为一级标题: \"{para_text_preview}...\"")
                    self.document_formatter._apply_font_to_runs(para, self.config['h1_font'], self.config['h1_size'],
                                                                set_color=apply_color)
                elif re_h2.match(text_to_check):
                    self._log(f"  > 文字识别为二级标题: \"{para_text_preview}...\"")
                    self.document_formatter._apply_font_to_runs(para, self.config['h2_font'], self.config['h2_size'],
                                                                set_color=apply_color)
                elif re_h3.match(text_to_check):
                    self._log(f"  > 文字识别为三级标题: \"{para_text_preview}...\"")
                    self.document_formatter._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'],
                                                                set_color=apply_color,
                                                                use_times_roman_for_ascii=self.config.get('body_use_times_roman', True))
                elif re_h4.match(text_to_check):
                    self._log(f"  > 文字识别为四级标题: \"{para_text_preview}...\"")
                    self.document_formatter._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'],
                                                                set_color=apply_color,
                                                                use_times_roman_for_ascii=self.config.get('body_use_times_roman', True))
                elif text_to_check:
                    self._log(f"  > 文字识别为正文: \"{para_text_preview}...\"")
                    self.document_formatter._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'],
                                                                set_color=apply_color,
                                                                use_times_roman_for_ascii=self.config.get('body_use_times_roman', True))

                # 确保图片或附件中的文字（即使是标题）也不缩进
                if re_h1.match(text_to_check) or re_h2.match(text_to_check) or re_h3.match(text_to_check) or re_h4.match(text_to_check):
                    self.document_formatter._apply_text_indent_and_align(para)
                
                block_idx += 1
                continue

            original_text, text_to_check = para.text, para.text.lstrip()
            text_to_check_stripped = para.text.strip()
            leading_space_count = len(original_text) - len(text_to_check)
            para_text_preview = text_to_check[:30].replace("\n", " ")

            spacing = para._p.get_or_add_pPr().get_or_add_spacing()
            spacing.set(qn('w:beforeAutospacing'), '0')
            spacing.set(qn('w:afterAutospacing'), '0')
            para.paragraph_format.space_before, para.paragraph_format.space_after = Pt(0), Pt(0)
            para.paragraph_format.line_spacing = Pt(self.config['line_spacing'])

            # 检查段落的大纲级别
            outline_level = self.document_formatter._get_outline_level(para)

            # 特殊处理：如果段落以"表"开头，强制识别为表格标题，不作为普通标题处理
            if para.text.strip().startswith("表"):
                self._log(f"段落 {current_block_num}: 检测到以'表'开头的段落，强制识别为表格标题")
                # 应用表格标题格式
                self.document_formatter._strip_leading_whitespace(para)
                config_font = self.config['table_caption_font']
                config_size = self.config['table_caption_size']
                config_bold = self.config.get('table_caption_bold', False)
                self.document_formatter._apply_font_to_runs(para, config_font, config_size,
                                                           set_color=apply_color, is_bold=config_bold)
                # 表格标题居中对齐且不缩进
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.document_formatter._apply_text_indent_and_align(para)
                # 应用表格标题大纲级别设置
                if 'table_caption_outline_level' in self.config:
                    outline_level_value = self.config['table_caption_outline_level']
                    if outline_level_value != '无' and outline_level_value != '':
                        try:
                            level = int(outline_level_value)
                            if 1 <= level <= 9:
                                self.document_formatter._set_outline_level(para, level)
                                self._log(f"  > 已设置表格标题的大纲级别为 {level}")
                        except (ValueError, TypeError):
                            pass  # 忽略无效值
                self.document_formatter._reset_pagination_properties(para)
                block_idx += 1
                continue

            # 检查段落的样式名称是否为标题样式
            is_heading_style = False
            try:
                style_name = getattr(para.style, 'name', '')
                # 检查样式名称是否为标题样式，如"标题1"、"标题2"、"Heading 1"等
                if style_name and (style_name.startswith("标题") or style_name.startswith("Heading")):
                    is_heading_style = True
                    # 如果是标题样式但没有大纲级别，设置默认大纲级别
                    if outline_level is None:
                        # 尝试从标题样式名称中提取级别数字
                        level_match = re.search(r'\d+', style_name)
                        if level_match:
                            level = int(level_match.group())
                            if 1 <= level <= 9:
                                outline_level = level - 1  # 转换为0-8范围
            except (AttributeError, ValueError):
                pass  # 忽略获取样式时可能出现的错误

            # 如果段落有大纲级别或是标题样式，则不进行首行缩进
            if outline_level is not None or is_heading_style:
                level = (outline_level or 0) + 1  # 大纲级别0-8对应标题级别1-9
                self._log(f"段落 {current_block_num}: 大纲级别 {level} 标题 - \"{para_text_preview}...\"")
                self.document_formatter._strip_leading_whitespace(para)

                # 根据大纲级别应用不同的字体和格式
                if level == 1:
                    # 检查配置中是否有h1_bold属性，如果没有则默认为False
                    h1_bold = self.config.get('h1_bold', False)
                    self.document_formatter._apply_font_to_runs(para, self.config['h1_font'], self.config['h1_size'],
                                                                set_color=apply_color, is_bold=h1_bold)
                    # 应用一级标题段前、段后间距（直接使用磅值）
                    h1_space_before_pts = float(self.config['h1_space_before'])
                    h1_space_after_pts = float(self.config['h1_space_after'])
                    para.paragraph_format.space_before = Pt(h1_space_before_pts)
                    para.paragraph_format.space_after = Pt(h1_space_after_pts)
                elif level == 2:
                    # 检查配置中是否有h2_bold属性，如果没有则默认为True
                    h2_bold = self.config.get('h2_bold', True)
                    self.document_formatter._apply_font_to_runs(para, self.config['h2_font'], self.config['h2_size'],
                                                                set_color=apply_color, is_bold=h2_bold)
                    # 应用二级标题段前、段后间距（直接使用磅值）
                    h2_space_before_pts = float(self.config['h2_space_before'])
                    h2_space_after_pts = float(self.config['h2_space_after'])
                    para.paragraph_format.space_before = Pt(h2_space_before_pts)
                    para.paragraph_format.space_after = Pt(h2_space_after_pts)
                elif level == 3:
                    # 检查配置中是否有h3_bold属性，如果没有则默认为False
                    h3_bold = self.config.get('h3_bold', False)
                    # 三级标题样式和间距设置
                    self.document_formatter._apply_font_to_runs(para, self.config['h3_font'], self.config['h3_size'],
                                                                set_color=apply_color, is_bold=h3_bold)
                    # 应用三级标题段前、段后间距（直接使用磅值）
                    h3_space_before_pts = float(self.config['h3_space_before'])
                    h3_space_after_pts = float(self.config['h3_space_after'])
                    para.paragraph_format.space_before = Pt(h3_space_before_pts)
                    para.paragraph_format.space_after = Pt(h3_space_after_pts)
                else:
                    # 4-9级标题使用正文字体和默认间距
                    self.document_formatter._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'],
                                                                set_color=apply_color,
                                                                use_times_roman_for_ascii=self.config.get('body_use_times_roman', True))
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)

                # 设置标题行间距
                spacing = para._p.get_or_add_pPr().get_or_add_spacing()
                spacing.set(qn('w:beforeAutospacing'), '0')
                spacing.set(qn('w:afterAutospacing'), '0')
                para.paragraph_format.line_spacing = Pt(self.config['line_spacing'])

                # 标题不缩进 - 确保所有标题（1-9级）都不缩进
                self.document_formatter._apply_text_indent_and_align(para)
                self.document_formatter._reset_pagination_properties(para)
                block_idx += 1
                continue

            # 检查是否为单独占一行的数字编号标题
            if not (outline_level is not None or is_heading_style):
                # 检查是否单独成行（排除段落中间的数字编号）
                lines = para.text.split('\n')
                if len(lines) == 1:
                    # 使用正则表达式精确匹配标题格式，避免将普通文本中的小数点误识别
                    text = para.text.strip()
                    
                    # 2级标题格式: "7.9 文本" 或 "7.9. 文本" - 数字后接数字的格式
                    h2_pattern = re.compile(r'^\d+[\.．]\d+(?:[\.．]\s*)?\s+[\u4e00-\u9fa5a-zA-Z]')
                    # 3级标题格式: "7.9.4 文本" 或 "7.9.4. 文本" - 三个数字的格式
                    h3_pattern = re.compile(r'^\d+[\.．]\d+[\.．]\d+(?:[\.．]\s*)?\s+[\u4e00-\u9fa5a-zA-Z]')
                    # 4级标题格式: "7.9.4.1 文本" 或 "7.9.4.1. 文本" - 四个数字的格式
                    h4_pattern = re.compile(r'^\d+[\.．]\d+[\.．]\d+[\.．]\d+(?:[\.．]\s*)?\s+[\u4e00-\u9fa5a-zA-Z]')
                    # 5级标题格式: "7.9.4.1.1 文本" 或 "7.9.4.1.1. 文本" - 五个或更多数字的格式
                    h5_pattern = re.compile(r'^\d+[\.．]\d+[\.．]\d+[\.．]\d+[\.．]\d+(?:[\.．]\d*)*(?:[\.．]\s*)?\s+[\u4e00-\u9fa5a-zA-Z]')
                    
                    if h2_pattern.match(text):
                        self._log(f"段落 {current_block_num}: 单独成行的2级数字编号标题 - \"{para_text_preview}...\"")
                        # 应用2级标题格式
                        self.document_formatter._strip_leading_whitespace(para)
                        h2_bold = self.config.get('h2_bold', True)
                        self.document_formatter._apply_font_to_runs(para, self.config['h2_font'], self.config['h2_size'],
                                                                    set_color=apply_color, is_bold=h2_bold)
                        # 应用二级标题段前、段后间距
                        h2_space_before_pts = float(self.config['h2_space_before'])
                        h2_space_after_pts = float(self.config['h2_space_after'])
                        para.paragraph_format.space_before = Pt(h2_space_before_pts)
                        para.paragraph_format.space_after = Pt(h2_space_after_pts)
                        # 设置标题行间距
                        spacing = para._p.get_or_add_pPr().get_or_add_spacing()
                        spacing.set(qn('w:beforeAutospacing'), '0')
                        spacing.set(qn('w:afterAutospacing'), '0')
                        para.paragraph_format.line_spacing = Pt(self.config['line_spacing'])
                        # 标题不缩进
                        self.document_formatter._apply_text_indent_and_align(para)
                        # 如果启用了大纲级别设置，则设置为2级
                        if self.config['set_outline']:
                            self.document_formatter._set_outline_level(para, 2)
                            self._log(f"  > 已设置为2级大纲级别")
                        # 标记此段落不需要缩进
                        setattr(para, '_has_no_indent', True)
                        self.document_formatter._reset_pagination_properties(para)
                        block_idx += 1
                        continue
                    elif h3_pattern.match(text):
                        self._log(f"段落 {current_block_num}: 单独成行的3级数字编号标题 - \"{para_text_preview}...\"")
                        # 应用3级标题格式
                        self.document_formatter._strip_leading_whitespace(para)
                        h3_bold = self.config.get('h3_bold', False)
                        self.document_formatter._apply_font_to_runs(para, self.config['h3_font'], self.config['h3_size'],
                                                                    set_color=apply_color, is_bold=h3_bold)
                        # 应用三级标题段前、段后间距
                        h3_space_before_pts = float(self.config['h3_space_before'])
                        h3_space_after_pts = float(self.config['h3_space_after'])
                        para.paragraph_format.space_before = Pt(h3_space_before_pts)
                        para.paragraph_format.space_after = Pt(h3_space_after_pts)
                        # 设置标题行间距
                        spacing = para._p.get_or_add_pPr().get_or_add_spacing()
                        spacing.set(qn('w:beforeAutospacing'), '0')
                        spacing.set(qn('w:afterAutospacing'), '0')
                        para.paragraph_format.line_spacing = Pt(self.config['line_spacing'])
                        # 标题不缩进
                        self.document_formatter._apply_text_indent_and_align(para)
                        # 如果启用了大纲级别设置，则设置为3级
                        if self.config['set_outline']:
                            self.document_formatter._set_outline_level(para, 3)
                            self._log(f"  > 已设置为3级大纲级别")
                        # 标记此段落不需要缩进
                        setattr(para, '_has_no_indent', True)
                        self.document_formatter._reset_pagination_properties(para)
                        block_idx += 1
                        continue
                    elif h4_pattern.match(text):
                        self._log(f"段落 {current_block_num}: 单独成行的4级数字编号标题 - \"{para_text_preview}...\"")
                        # 应用4级标题格式（使用正文字体）
                        self.document_formatter._strip_leading_whitespace(para)
                        # 4级标题使用正文字体
                        self.document_formatter._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'],
                                                                    set_color=apply_color,
                                                                    use_times_roman_for_ascii=self.config.get('body_use_times_roman', True))
                        # 设置默认间距
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                        # 设置标题行间距
                        spacing = para._p.get_or_add_pPr().get_or_add_spacing()
                        spacing.set(qn('w:beforeAutospacing'), '0')
                        spacing.set(qn('w:afterAutospacing'), '0')
                        para.paragraph_format.line_spacing = Pt(self.config['line_spacing'])
                        # 标题不缩进
                        self.document_formatter._apply_text_indent_and_align(para)
                        # 如果启用了大纲级别设置，则设置为4级
                        if self.config['set_outline']:
                            self.document_formatter._set_outline_level(para, 4)
                            self._log(f"  > 已设置为4级大纲级别")
                        # 标记此段落不需要缩进
                        setattr(para, '_has_no_indent', True)
                        self.document_formatter._reset_pagination_properties(para)
                        block_idx += 1
                        continue
                    elif h5_pattern.match(text):
                        self._log(f"段落 {current_block_num}: 单独成行的5级数字编号标题 - \"{para_text_preview}...\"")
                        # 应用5级标题格式（使用正文字体）
                        self.document_formatter._strip_leading_whitespace(para)
                        # 5级标题使用正文字体
                        self.document_formatter._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'],
                                                                    set_color=apply_color,
                                                                    use_times_roman_for_ascii=self.config.get('body_use_times_roman', True))
                        # 设置默认间距
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                        # 设置标题行间距
                        spacing = para._p.get_or_add_pPr().get_or_add_spacing()
                        spacing.set(qn('w:beforeAutospacing'), '0')
                        spacing.set(qn('w:afterAutospacing'), '0')
                        para.paragraph_format.line_spacing = Pt(self.config['line_spacing'])
                        # 标题不缩进
                        self.document_formatter._apply_text_indent_and_align(para)
                        # 如果启用了大纲级别设置，则设置为5级
                        if self.config['set_outline']:
                            self.document_formatter._set_outline_level(para, 5)
                            self._log(f"  > 已设置为5级大纲级别")
                        # 标记此段落不需要缩进
                        setattr(para, '_has_no_indent', True)
                        self.document_formatter._reset_pagination_properties(para)
                        block_idx += 1
                        continue
            # 取消自动识别"一、"、"（一）"、"1."、"(1)"等常规标题的功能
            # 直接将这些段落作为正文处理
            if re_h1.match(text_to_check) or re_h2.match(text_to_check) or re_h3.match(text_to_check) or re_h4.match(
                    text_to_check):
                self._log(
                    f"段落 {current_block_num}: 常规标题格式文本 - \"{para_text_preview}...\" (已禁用自动识别，按正文处理)")

            # 所有段落都按正文处理
            self._log(f"段落 {current_block_num}: 正文 - \"{para_text_preview}...\"")
            self.document_formatter._strip_leading_whitespace(para)
            self.document_formatter._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'],
                                                        set_color=apply_color,
                                                        use_times_roman_for_ascii=self.config.get('body_use_times_roman', True))
            # 正文需要首行缩进
            self.document_formatter._apply_body_text_indent_and_align(para)
            self.document_formatter._reset_pagination_properties(para)

            block_idx += 1

        self.page_setup._apply_page_setup(doc, is_from_txt=is_from_txt)
        self._log("正在保存最终文档...")
        doc.save(output_path)

    def _cleanup_temp_files(self):
        self.file_processor._cleanup_temp_files()

    def quit_com_app(self):
        self.file_processor.quit_com_app()