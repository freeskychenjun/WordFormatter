import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, Menu
import json
import os
import re
import logging
import shutil
import win32com.client
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from tkinterdnd2 import DND_FILES, TkinterDnD

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class WordProcessor:
    def __init__(self, config, log_callback=None):
        self.config = config
        self.temp_files = []
        self.log_callback = log_callback
        self.com_app = None

    def _log(self, message):
        if self.log_callback: self.log_callback(message)

    def _cleanup_temp_files(self):
        self._log("正在清理本轮临时文件...")
        for f in self.temp_files:
            try:
                if os.path.exists(f):
                    os.remove(f)
                    self._log(f"  > 临时文件 {os.path.basename(f)} 已删除")
            except OSError as e:
                self._log(f"  > 警告：删除临时文件 {f} 失败: {e}")
        self.temp_files.clear()

    def _get_wps_app(self):
        if self.com_app is None:
            self._log("首次需要，正在启动WPS/Word应用...")
            try:
                self.com_app = win32com.client.Dispatch('KWPS.Application')
                self._log("  > 已成功连接到WPS。")
            except Exception:
                try:
                    self.com_app = win32com.client.Dispatch('Word.Application')
                    self._log("  > 已成功连接到Word。")
                except Exception as e:
                    raise RuntimeError(f"未能启动WPS或Word，请确保已安装。错误: {e}")
            self.com_app.Visible = False
        return self.com_app
        
    def quit_com_app(self):
        if self.com_app:
            self._log("所有任务完成，正在关闭WPS/Word应用...")
            self.com_app.Quit()
            self.com_app = None
            self._log("  > 应用已关闭。")



    def convert_to_docx(self, input_path):
        file_ext = os.path.splitext(input_path)[1].lower()
        is_from_txt = (file_ext == '.txt')
        
        # 使用系统临时目录来避免权限问题
        temp_dir = tempfile.gettempdir()
        base_name = os.path.splitext(os.path.basename(input_path))[0]

        # 清理文件名中的特殊字符，避免在Windows系统中出现问题
        # 移除或替换可能引起问题的字符
        cleaned_base_name = re.sub(r'[<>:"/\\|?*\x00-\x1F]', '_', base_name)
        # 限制文件名长度，避免超过系统限制
        cleaned_base_name = cleaned_base_name[:100] if len(cleaned_base_name) > 100 else cleaned_base_name
        # 移除末尾的点和空格（Windows不允许）
        cleaned_base_name = cleaned_base_name.rstrip('. ')
        
        # 添加调试信息
        self._log(f"  > 原始文件路径: {input_path}")
        self._log(f"  > 系统临时目录: {temp_dir}")
        self._log(f"  > 清理后的文件名: {cleaned_base_name}")

        if file_ext == '.docx':
            self._log("检测到 .docx 文件，正在创建安全的处理副本...")
            temp_docx_path = os.path.join(temp_dir, f"~temp_copy_{cleaned_base_name}.docx")
            self._log(f"  > 临时文件路径: {temp_docx_path}")
            
            # 检查目标路径是否存在同名文件，如果存在则删除
            if os.path.exists(temp_docx_path):
                self._log(f"  > 检测到同名临时文件，正在删除: {temp_docx_path}")
                try:
                    os.remove(temp_docx_path)
                except Exception as e:
                    self._log(f"  > 删除同名临时文件失败: {e}")
            
            try:
                shutil.copy2(input_path, temp_docx_path)
                self.temp_files.append(temp_docx_path)
                self._log(f"  > 副本创建成功: {os.path.basename(temp_docx_path)}")
                
                # 验证文件是否创建成功
                if os.path.exists(temp_docx_path):
                    self._log(f"  > 验证文件存在: {temp_docx_path}")
                else:
                    self._log(f"  > 警告：文件创建后不存在: {temp_docx_path}")
                
                return temp_docx_path, is_from_txt
            except Exception as e:
                self._log(f"  > 创建副本失败: {e}")
                # 如果复制失败，尝试使用不同的临时文件名
                for i in range(10):
                    alt_temp_docx_path = os.path.join(temp_dir, f"~temp_copy_{cleaned_base_name}_{i}.docx")
                    self._log(f"  > 尝试备用路径: {alt_temp_docx_path}")
                    if not os.path.exists(alt_temp_docx_path):
                        try:
                            shutil.copy2(input_path, alt_temp_docx_path)
                            self.temp_files.append(alt_temp_docx_path)
                            self._log(f"  > 使用备用名称创建副本成功: {os.path.basename(alt_temp_docx_path)}")
                            return alt_temp_docx_path, is_from_txt
                        except Exception as e2:
                            self._log(f"  > 备用名称创建副本也失败: {e2}")
                            continue
                # 如果所有尝试都失败，抛出异常
                raise RuntimeError(f"无法创建文件副本: {e}")

        temp_docx_path = os.path.join(temp_dir, f"~temp_converted_{cleaned_base_name}.docx")
        self.temp_files.append(temp_docx_path)
        self._log(f"  > 转换文件路径: {temp_docx_path}")

        if file_ext == '.txt':
            self._log("检测到 .txt 文件，正在创建 .docx...")
            doc = Document()
            try:
                with open(input_path, 'r', encoding='utf-8') as f:
                    for line in f: doc.add_paragraph(line.strip())
                self._log("  > 已使用 UTF-8 编码读取TXT文件。")
            except UnicodeDecodeError:
                self._log("  > UTF-8读取失败，尝试使用 GBK 编码...")
                with open(input_path, 'r', encoding='gbk') as f:
                    for line in f: doc.add_paragraph(line.strip())
                self._log("  > 已成功使用 GBK 编码读取TXT文件。")
            doc.save(temp_docx_path)
            self._log("TXT转换完成。")
            
            # 验证文件是否创建成功
            if os.path.exists(temp_docx_path):
                self._log(f"  > 验证文件存在: {temp_docx_path}")
            else:
                self._log(f"  > 警告：文件创建后不存在: {temp_docx_path}")
                
            return temp_docx_path, is_from_txt
        elif file_ext in ['.wps', '.doc']:
            self._log(f"正在转换 {file_ext} 文件为 .docx...")
            app = self._get_wps_app()
            doc_com = app.Documents.Open(os.path.abspath(input_path), ReadOnly=1)
            doc_com.SaveAs2(os.path.abspath(temp_docx_path), FileFormat=12)
            doc_com.Close()
            self._log("文件格式转换完成。")
            
            # 验证文件是否创建成功
            if os.path.exists(temp_docx_path):
                self._log(f"  > 验证文件存在: {temp_docx_path}")
            else:
                self._log(f"  > 警告：文件创建后不存在: {temp_docx_path}")
                
            return temp_docx_path, is_from_txt
        
        raise ValueError(f"不支持的文件格式: {file_ext}")

    def _preprocess_com_tasks(self, docx_path):
        self._log("正在对副本执行预处理（接受所有修订、转换自动编号）...")
        app = self._get_wps_app()
        try:
            doc_com = app.Documents.Open(os.path.abspath(docx_path))
            
            doc_com.TrackRevisions = False
            self._log("  > 已关闭修订追踪。")
            
            if doc_com.Revisions.Count > 0:
                doc_com.AcceptAllRevisions()
                self._log("  > 已接受文档副本中的所有修订。")
            
            doc_com.Content.ListFormat.ConvertNumbersToText()
            self._log("  > 已将副本中的自动编号转换为文本。")
            
            if doc_com.Revisions.Count > 0:
                doc_com.AcceptAllRevisions()
                self._log("  > 已接受编号转换产生的修订。")
            
            doc_com.TrackRevisions = False
            
            # 确保文档被正确保存
            doc_com.Save()
            doc_com.Close()
            self._log("预处理完成。")
            
            # 添加延迟以确保文件操作完成
            import time
            time.sleep(0.5)
            
            # 验证文件是否存在
            if os.path.exists(docx_path):
                self._log(f"  > 验证预处理后文件存在: {docx_path}")
            else:
                self._log(f"  > 警告：预处理后文件不存在: {docx_path}")
        except Exception as e:
            self._log(f"警告：执行预处理任务时出错: {e}")

    def _create_page_number(self, paragraph, text):
        font_name = self.config['page_number_font']
        font_size = self.config['page_number_size']
        self._set_run_font(paragraph.add_run('— '), font_name, font_size, set_color=True)
        run_field = paragraph.add_run()
        self._set_run_font(run_field, font_name, font_size, set_color=True)
        fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = text
        fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
        run_field._r.extend([fldChar1, instrText, fldChar2])
        self._set_run_font(paragraph.add_run(' —'), font_name, font_size, set_color=True)

    def _apply_page_setup(self, doc, is_from_txt=False):
        self._log("正在应用页面边距设置...")
        
        # 判断是否需要强制设置A4纸
        # 逻辑：如果是纯文本来源（包括直接输入）则设置为A4
        should_set_a4 = is_from_txt

        for section in doc.sections:
            section.top_margin = Cm(self.config['margin_top'])
            section.bottom_margin = Cm(self.config['margin_bottom'])
            section.left_margin = Cm(self.config['margin_left'])
            section.right_margin = Cm(self.config['margin_right'])

            # 设置纸张大小为A4 (仅在需要时)
            if should_set_a4:
                section.page_width = Cm(21)
                section.page_height = Cm(29.7)
        
        if should_set_a4:
            self._log("  > 已将页面大小设置为 A4。")

    def _set_run_font(self, run, font_name, size_pt, set_color=False, is_bold=False):
        """设置单个run的字体属性
        
        Args:
            run: 要设置的run对象
            font_name: 字体名称
            size_pt: 字号（磅值）
            set_color: 是否设置颜色为黑色
            is_bold: 是否设置为粗体（新增属性）
        """
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.bold = is_bold  # 设置粗体属性

        if set_color: run.font.color.rgb = RGBColor(0, 0, 0)
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)



    def _apply_font_to_runs(self, para, font_name, size_pt, set_color=False, is_bold=False):
        """应用字体设置到段落的所有runs
        
        Args:
            para: 段落对象
            font_name: 字体名称
            size_pt: 字号（磅值）
            set_color: 是否设置颜色为黑色
            is_bold: 是否设置为粗体（新增属性）
        """
        for run in para.runs: self._set_run_font(run, font_name, size_pt, set_color=set_color, is_bold=is_bold)

    def _get_paragraph_font_info(self, para):
        """获取段落主要字体和字号信息"""
        if not para.runs:
            return None, None
        
        # 获取第一个非空run的字体信息
        for run in para.runs:
            if run.text.strip():
                font_name = run.font.name
                font_size = run.font.size.pt if run.font.size else None
                return font_name, font_size
        return None, None

    def _strip_leading_whitespace(self, para):
        if not para.runs: return
        while para.runs and not para.runs[0].text.strip():
            p = para._p
            p.remove(para.runs[0]._r)
        if not para.runs: return
        first_run = para.runs[0]
        original_text = first_run.text
        stripped_text = original_text.lstrip()
        if original_text != stripped_text:
            first_run.text = stripped_text
            self._log("  > 已移除段落前的多余空格。")
    
    def _reset_pagination_properties(self, para):
        para.paragraph_format.widow_control = False
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.keep_lines_together = False
        para.paragraph_format.page_break_before = False
        para.paragraph_format.keep_together = False

    def _get_outline_level(self, para):
        """
        读取段落的当前大纲级别
        返回: 0-8 表示级别1-9，None 表示未设置
        """
        pPr = para._p.get_or_add_pPr()
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is not None:
            val = outlineLvl.get(qn('w:val'))
            if val is not None:
                return int(val)
        return None

    def _set_outline_level(self, para, level):
        """
        直接设置段落的大纲级别，不通过样式，不影响字体字号等格式
        level: 1-9 的整数，表示大纲级别
        返回: 原有的大纲级别 (0-8) 或 None
        """
        if level < 1 or level > 9:
            self._log(f"  > 警告：大纲级别 {level} 超出范围 (1-9)，已跳过设置")
            return None
        
        # 读取原有大纲级别
        original_level = self._get_outline_level(para)
        
        # 设置新的大纲级别 (Word内部用0-8表示1-9级)
        pPr = para._p.get_or_add_pPr()
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is None:
            outlineLvl = OxmlElement('w:outlineLvl')
            pPr.append(outlineLvl)
        outlineLvl.set(qn('w:val'), str(level - 1))
        
        return original_level

    def _apply_text_indent_and_align(self, para):
        # 标题不缩进
        para.paragraph_format.first_line_indent = None
        para.paragraph_format.left_indent = Pt(0)
        # para.paragraph_format.right_indent = Cm(self.config['right_indent_cm'])
        # 不设置首行缩进
        # para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # 显式设置首行缩进字符数为0，确保所有有大纲级别的内容都不会应用首行缩进
        # 同时设置leftChars为0，确保不产生任何左缩进
        ind = para._p.get_or_add_pPr().get_or_add_ind()
        ind.set(qn("w:firstLineChars"), "0")
        ind.set(qn("w:leftChars"), "0")

    def _apply_body_text_indent_and_align(self, para):
        # 正文首行缩进2字符
        # para.paragraph_format.left_indent = Cm(self.config['left_indent_cm'])
        # para.paragraph_format.right_indent = Cm(self.config['right_indent_cm'])
        # 设置首行缩进2字符（200表示2个字符）
        ind = para._p.get_or_add_pPr().get_or_add_ind()
        ind.set(qn("w:firstLineChars"), "200")
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _iter_block_items(self, parent):
        parent_elm = parent.element.body if isinstance(parent, _Document) else parent._tc
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P): yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl): yield Table(child, parent)
    
    def _find_title_and_subtitle_paragraphs(self, doc, is_from_txt, start_index=0):
        """
        查找题目和副标题段落的索引范围




        返回: (title_indices, subtitle_indices)
        title_indices: 题目行的索引列表
        subtitle_indices: 副标题行的索引列表
        """
        ch_num = r'[一二三四五六七八九十百千万零]+'
        re_h1 = re.compile(r'^' + ch_num + r'\s*、')
        re_h2 = re.compile(r'^[（\(]' + ch_num + r'[）\)]')

        all_blocks = list(self._iter_block_items(doc))
        
        # 查找首个标题行
        first_title_idx = -1
        
        if is_from_txt:
            self._log("文档源自 TXT，采用智能规则查找题目...")
            for idx in range(start_index, len(all_blocks)):
                block = all_blocks[idx]
                if isinstance(block, Paragraph) and block.text.strip():
                    text_to_check = block.text.strip()
                    if re_h1.match(text_to_check) or re_h2.match(text_to_check):
                        self._log(f"  > 首个非空行 (块 {idx + 1}) 符合标题格式，认定本文档无独立题目。")
                        return [], []
                    else:
                        self._log(f"  > 在块 {idx + 1} 发现首个非空段落，认定为题目首行。")
                        first_title_idx = idx
                        break
        else:
            self._log("正在预扫描以确定居中题目位置...")
            for idx in range(start_index, len(all_blocks)):
                block = all_blocks[idx]
                if not isinstance(block, Paragraph) or not block.text.strip(): 
                    continue
                para = block
                text_to_check = para.text.lstrip()
                if re_h1.match(text_to_check) or re_h2.match(text_to_check):
                    self._log("  > 发现一级/二级标题，在此之前未找到居中题目。")
                    return [], []
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    self._log(f"  > 在块 {idx + 1} 发现潜在题目首行。")
                    first_title_idx = idx
                    break
        
        if first_title_idx == -1:
            self._log("  > 扫描结束，未能找到题目。")
            return [], []
        
        # 获取首个标题行的字体字号信息
        first_title_para = all_blocks[first_title_idx]
        title_font, title_size = self._get_paragraph_font_info(first_title_para)
        
        # 向下查找连续的标题行
        title_indices = [first_title_idx]
        idx = first_title_idx + 1
        
        while idx < len(all_blocks):
            block = all_blocks[idx]
            if not isinstance(block, Paragraph):
                break
            
            para = block
            text = para.text.strip()
            
            # 遇到空行，停止标题识别
            if not text:
                self._log(f"  > 在块 {idx + 1} 遇到空行，标题识别结束。")
                break
            
            # 检查是否居中
            if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                break
            
            # 检查字体字号是否与首行相同
            para_font, para_size = self._get_paragraph_font_info(para)
            if para_font == title_font and para_size == title_size:
                self._log(f"  > 块 {idx + 1} 也是标题行（居中且字体字号相同）。")
                title_indices.append(idx)
                idx += 1
            else:
                # 字体字号不同，可能是副标题的开始
                break
        
        self._log(f"  > 共识别到 {len(title_indices)} 行标题。")
        
        # 查找副标题
        subtitle_indices = []
        subtitle_start_idx = idx
        
        # 跳过空行
        while subtitle_start_idx < len(all_blocks):
            block = all_blocks[subtitle_start_idx]
            if isinstance(block, Paragraph) and block.text.strip():
                break
            if isinstance(block, Paragraph):
                subtitle_start_idx += 1
            else:
                # 遇到非段落（如表格），停止
                break
        
        # 检查是否有副标题
        if subtitle_start_idx < len(all_blocks):
            block = all_blocks[subtitle_start_idx]
            if isinstance(block, Paragraph):
                para = block
                text = para.text.strip()
                
                # 副标题必须居中
                if text and para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    # 检查字体字号是否与标题不同
                    para_font, para_size = self._get_paragraph_font_info(para)
                    if para_font != title_font or para_size != title_size:
                        self._log(f"  > 在块 {subtitle_start_idx + 1} 发现副标题首行（居中且字体字号与标题不同）。")
                        subtitle_indices.append(subtitle_start_idx)
                        
                        # 查找连续的副标题行
                        subtitle_font, subtitle_size = para_font, para_size
                        idx = subtitle_start_idx + 1
                        
                        while idx < len(all_blocks):
                            block = all_blocks[idx]
                            if not isinstance(block, Paragraph):
                                break
                            
                            para = block
                            text = para.text.strip()
                            
                            # 遇到空行，停止副标题识别
                            if not text:
                                self._log(f"  > 在块 {idx + 1} 遇到空行，副标题识别结束。")
                                break
                            
                            # 检查是否居中
                            if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                                break
                            
                            # 检查字体字号是否与副标题首行相同
                            para_font, para_size = self._get_paragraph_font_info(para)
                            if para_font == subtitle_font and para_size == subtitle_size:
                                self._log(f"  > 块 {idx + 1} 也是副标题行（居中且字体字号相同）。")
                                subtitle_indices.append(idx)
                                idx += 1
                            else:
                                break
                        
                        self._log(f"  > 共识别到 {len(subtitle_indices)} 行副标题。")
        
        return title_indices, subtitle_indices

    def format_document(self, input_path, output_path):
        processing_path, is_from_txt = self.convert_to_docx(input_path)
        self._log(f"  > 处理路径: {processing_path}")
        
        # 检查临时文件是否存在
        if not os.path.exists(processing_path):
            self._log(f"  > 错误：临时文件不存在: {processing_path}")
            raise FileNotFoundError(f"临时文件不存在: {processing_path}")
        
        if not is_from_txt: 
            self._preprocess_com_tasks(processing_path)
            # 预处理完成后，再次检查文件是否存在
            if not os.path.exists(processing_path):
                self._log(f"  > 警告：预处理后文件不存在: {processing_path}")
                # 尝试重新创建文件
                self._log(f"  > 尝试重新创建文件...")
                processing_path, _ = self.convert_to_docx(input_path)
                if not os.path.exists(processing_path):
                    self._log(f"  > 重新创建文件也失败: {processing_path}")
                    raise FileNotFoundError(f"预处理后文件不存在且重新创建失败: {processing_path}")
                else:
                    self._log(f"  > 重新创建文件成功: {processing_path}")
            else:
                self._log(f"  > 验证预处理后文件仍存在: {processing_path}")
        
        # 再次检查文件是否存在
        if not os.path.exists(processing_path):
            self._log(f"  > 错误：预处理后文件不存在: {processing_path}")
            raise FileNotFoundError(f"预处理后文件不存在: {processing_path}")
        
        # 尝试打开文档，如果失败则尝试重新创建
        doc = None
        try:
            doc = Document(processing_path)
        except Exception as e:
            self._log(f"  > 首次打开文档失败: {e}")
            # 尝试使用绝对路径
            abs_path = os.path.abspath(processing_path)
            self._log(f"  > 尝试使用绝对路径: {abs_path}")
            try:
                doc = Document(abs_path)
            except Exception as e2:
                self._log(f"  > 使用绝对路径也失败: {e2}")
                # 如果还是失败，尝试重新创建文件
                self._log(f"  > 尝试重新创建文件...")
                processing_path, _ = self.convert_to_docx(input_path)
                if os.path.exists(processing_path):
                    self._log(f"  > 重新创建文件成功，再次尝试打开...")
                    try:
                        doc = Document(processing_path)
                        self._log(f"  > 重新打开文档成功")
                    except Exception as e3:
                        self._log(f"  > 重新打开文档也失败: {e3}")
                        raise
                else:
                    self._log(f"  > 重新创建文件失败: {processing_path}")
                    raise FileNotFoundError(f"重新创建文件失败: {processing_path}")
        
        if doc is None:
            raise RuntimeError("无法打开文档")
        
        all_blocks = list(self._iter_block_items(doc))
        processed_indices = set()
        
        apply_color = not is_from_txt

        if not is_from_txt:
            self._log("正在扫描图表标题...")
            for idx, block in enumerate(all_blocks):
                is_pic_para = isinstance(block, Paragraph) and ('<w:drawing>' in block._p.xml or '<w:pict>' in block._p.xml)
                is_table = isinstance(block, Table)
                
                if not (is_pic_para or is_table): continue
                
                for direction in [-1, 1]:
                    caption_found = False
                    for i in range(idx + direction, -1 if direction == -1 else len(all_blocks), direction):
                        if i in processed_indices: continue
                        potential_caption = all_blocks[i]
                        if not isinstance(potential_caption, Paragraph): break 
                        text = potential_caption.text.strip()
                        if text: 
                            if potential_caption.alignment == WD_ALIGN_PARAGRAPH.CENTER and (text.startswith("图") or text.startswith("表")):
                                detected_type = "图" if text.startswith("图") else "表"
                                self._log(f"  > 发现 {detected_type} 的标题: \"{text[:30]}...\" (在段落 {i+1})")
                                config_font_key = f'{("figure" if detected_type == "图" else "table")}_caption_font'
                                config_size_key = f'{("figure" if detected_type == "图" else "table")}_caption_size'
                                config_font = self.config[config_font_key]
                                config_size = self.config[config_size_key]
                                config_bold_key = f'{("figure" if detected_type == "图" else "table")}_caption_bold'
                                config_bold = self.config.get(config_bold_key, False)
                                self._apply_font_to_runs(potential_caption, config_font, config_size, set_color=apply_color, is_bold=config_bold)
                                # 表格/图表标题不缩进，确保完全没有任何缩进
                                potential_caption.paragraph_format.first_line_indent = None
                                potential_caption.paragraph_format.left_indent = Pt(0)
                                potential_caption.paragraph_format.right_indent = Pt(0)
                                # 确保完全清除任何可能的缩进设置
                                ind = potential_caption._p.get_or_add_pPr().get_or_add_ind()
                                ind.set(qn("w:firstLineChars"), "0")
                                ind.set(qn("w:leftChars"), "0")
                                # 移除任何可能存在的其他缩进相关属性
                                if hasattr(ind, 'get_or_add_firstLine'):
                                    try:
                                        ind.remove(ind.get_or_add_firstLine())
                                    except:
                                        pass
                                # 确保应用正确的缩进设置
                                if qn("w:firstLine") in [child.tag for child in ind]:
                                    for child in ind:
                                        if child.tag == qn("w:firstLine"):
                                            ind.remove(child)
                                            break
                                
                                # 应用大纲级别设置
                                outline_level_key = f'{("figure" if detected_type == "图" else "table")}_caption_outline_level'
                                if outline_level_key in self.config:
                                    outline_level_value = self.config[outline_level_key]
                                    if outline_level_value != '无' and outline_level_value != '':
                                        try:
                                            level = int(outline_level_value)
                                            if 1 <= level <= 9:
                                                self._set_outline_level(potential_caption, level)
                                                self._log(f"  > 已设置 {detected_type} 标题的大纲级别为 {level}")
                                        except (ValueError, TypeError):
                                            pass  # 忽略无效值
                                
                                processed_indices.add(i)
                                caption_found = True
                            break 
                    if caption_found: break 

        # 查找主标题和副标题
        title_indices, subtitle_indices = self._find_title_and_subtitle_paragraphs(doc, is_from_txt)
        
        # 将标题和副标题索引加入已处理集合
        for idx in title_indices:
            processed_indices.add(idx)
        for idx in subtitle_indices:
            processed_indices.add(idx)

        self._log("预扫描完成，开始逐段格式化...")
        if self.config['set_outline']:
            self._log("【大纲级别设置已启用】")
        else:


            self._log("【大纲级别设置已禁用】")
            
        re_h1 = re.compile(r'^[一二三四五六七八九十百千万零]+\s*、')
        re_h2 = re.compile(r'^[（\(][一二三四五六七八九十百千万零]+[）\)]')
        re_h3 = re.compile(r'^\d+\s*[\.．]')
        re_h4 = re.compile(r'^[（\(]\d+[）\)]')
        re_attachment = re.compile(r'^附件\s*(\d+|[一二三四五六七八九十百千万零]+)?\s*[:：]?$')

        # 格式化主标题
        if title_indices:
            self._log(f"\n开始格式化主标题（共 {len(title_indices)} 行）...")
            for idx in title_indices:
                para = all_blocks[idx]
                self._log(f"段落 {idx + 1}: 主标题行 - \"{para.text[:30]}...\"")
                self._strip_leading_whitespace(para)
                self._apply_font_to_runs(para, self.config['h1_font'], self.config['h1_size'], set_color=apply_color)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = None
                para.paragraph_format.left_indent = Pt(0)
                
                # 设置标题行间距
                spacing = para._p.get_or_add_pPr().get_or_add_spacing()
                spacing.set(qn('w:beforeAutospacing'), '0')
                spacing.set(qn('w:afterAutospacing'), '0')
                # 根据配置设置一级标题段前、段后间距（直接使用磅值）
                h1_space_before_pts = float(self.config['h1_space_before'])
                h1_space_after_pts = float(self.config['h1_space_after'])
                para.paragraph_format.space_before = Pt(h1_space_before_pts)
                para.paragraph_format.space_after = Pt(h1_space_after_pts)
                para.paragraph_format.line_spacing = Pt(self.config['line_spacing'])
                
                self._reset_pagination_properties(para)
        
        # 格式化副标题
        if subtitle_indices:
            self._log(f"\n开始格式化副标题（共 {len(subtitle_indices)} 行）...")
            for idx in subtitle_indices:
                para = all_blocks[idx]
                self._log(f"段落 {idx + 1}: 副标题行 - \"{para.text[:30]}...\"")
                self._strip_leading_whitespace(para)
                self._apply_font_to_runs(para, self.config['h2_font'], self.config['h2_size'], set_color=apply_color)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = None
                para.paragraph_format.left_indent = Pt(0)
                
                # 设置副标题行间距
                spacing = para._p.get_or_add_pPr().get_or_add_spacing()
                spacing.set(qn('w:beforeAutospacing'), '0')
                spacing.set(qn('w:afterAutospacing'), '0')
                # 根据配置设置二级标题段前、段后间距（直接使用磅值）
                h2_space_before_pts = float(self.config['h2_space_before'])
                h2_space_after_pts = float(self.config['h2_space_after'])
                para.paragraph_format.space_before = Pt(h2_space_before_pts)
                para.paragraph_format.space_after = Pt(h2_space_after_pts)
                para.paragraph_format.line_spacing = Pt(self.config['line_spacing'])
                
                self._reset_pagination_properties(para)

        block_idx = 0
        while block_idx < len(all_blocks):
            block = all_blocks[block_idx]
            
            if block_idx in processed_indices:
                if block_idx not in title_indices and block_idx not in subtitle_indices:
                    self._log(f"块 {block_idx + 1}: 已作为图表/附件标题处理 - 跳过")
                block_idx += 1
                continue

            current_block_num = block_idx + 1
            if isinstance(block, Table): 
                self._log(f"块 {current_block_num}: 表格 - 检查内部标题")
                # 检查表格内部第一行是否为标题
                table = block
                if len(table.rows) > 0:
                    first_row = table.rows[0]
                    # 检查第一行的所有单元格
                    for cell in first_row.cells:
                        if len(cell.paragraphs) > 0:
                            for para in cell.paragraphs:
                                text = para.text.strip()
                                if text and para.alignment == WD_ALIGN_PARAGRAPH.CENTER and text.startswith("表"):
                                    self._log(f"  > 发现表格内部标题: \"{text[:30]}...\"")
                                    config_font = self.config['table_caption_font']
                                    config_size = self.config['table_caption_size']
                                    config_bold = self.config.get('table_caption_bold', False)
                                    self._apply_font_to_runs(para, config_font, config_size, set_color=apply_color, is_bold=config_bold)
                                    # 表格标题不缩进，确保完全没有任何缩进
                                    para.paragraph_format.first_line_indent = None
                                    para.paragraph_format.left_indent = Pt(0)
                                    para.paragraph_format.right_indent = Pt(0)
                                    # 确保完全清除任何可能的缩进设置
                                    ind = para._p.get_or_add_pPr().get_or_add_ind()
                                    ind.set(qn("w:firstLineChars"), "0")
                                    ind.set(qn("w:leftChars"), "0")
                                    # 移除任何可能存在的其他缩进相关属性
                                    if hasattr(ind, 'get_or_add_firstLine'):
                                        try:
                                            ind.remove(ind.get_or_add_firstLine())
                                        except:
                                            pass
                                    # 确保应用正确的缩进设置
                                    if qn("w:firstLine") in [child.tag for child in ind]:
                                        for child in ind:
                                            if child.tag == qn("w:firstLine"):
                                                ind.remove(child)
                                                break
                                    
                                    # 应用表格标题大纲级别设置
                                    if 'table_caption_outline_level' in self.config:
                                        outline_level_value = self.config['table_caption_outline_level']
                                        if outline_level_value != '无' and outline_level_value != '':
                                            try:
                                                level = int(outline_level_value)
                                                if 1 <= level <= 9:
                                                    self._set_outline_level(para, level)
                                                    self._log(f"  > 已设置表格内部标题的大纲级别为 {level}")
                                            except (ValueError, TypeError):
                                                pass  # 忽略无效值
                                    
                                    break
                block_idx += 1
                continue
            
            para = block
            if not para.text.strip(): 
                self._log(f"段落 {current_block_num}: 空白 - 跳过"); block_idx += 1; continue
            
            is_pic = '<w:drawing>' in para._p.xml or '<w:pict>' in para._p.xml
            is_embedded_obj = '<w:object>' in para._p.xml
            if is_pic or is_embedded_obj:
                log_msg = "图片" if is_pic else "附件"
                self._log(f"段落 {current_block_num}: {log_msg} - 仅格式化文字")
                
                text_to_check = para.text.lstrip()
                para_text_preview = text_to_check[:30].replace("\n", " ")

                if re_h1.match(text_to_check):
                    self._log(f"  > 文字识别为一级标题: \"{para_text_preview}...\"")
                    self._apply_font_to_runs(para, self.config['h1_font'], self.config['h1_size'], set_color=apply_color)
                elif re_h2.match(text_to_check):
                    self._log(f"  > 文字识别为二级标题: \"{para_text_preview}...\"")
                    self._apply_font_to_runs(para, self.config['h2_font'], self.config['h2_size'], set_color=apply_color)
                elif re_h3.match(text_to_check):
                    self._log(f"  > 文字识别为三级标题: \"{para_text_preview}...\"")
                    self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                elif re_h4.match(text_to_check):
                    self._log(f"  > 文字识别为四级标题: \"{para_text_preview}...\"")
                    self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                elif text_to_check:
                    self._log(f"  > 文字识别为正文: \"{para_text_preview}...\"")


                    self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)

                block_idx += 1
                continue

            original_text, text_to_check = para.text, para.text.lstrip()
            text_to_check_stripped = para.text.strip()
            leading_space_count = len(original_text) - len(text_to_check)
            para_text_preview = text_to_check[:30].replace("\n", " ")
            
            spacing = para._p.get_or_add_pPr().get_or_add_spacing()
            spacing.set(qn('w:beforeAutospacing'), '0'); spacing.set(qn('w:afterAutospacing'), '0')
            para.paragraph_format.space_before, para.paragraph_format.space_after = Pt(0), Pt(0)
            para.paragraph_format.line_spacing = Pt(self.config['line_spacing'])

            # 检查段落的大纲级别
            outline_level = self._get_outline_level(para)
            
            # 检查段落的样式名称是否为标题样式
            is_heading_style = False
            try:
                style_name = para.style.name
                # 检查样式名称是否为标题样式，如"标题1"、"标题2"、"Heading 1"等
                if style_name.startswith("标题") or style_name.startswith("Heading"):
                    is_heading_style = True
                    # 如果是标题样式但没有大纲级别，设置默认大纲级别
                    if outline_level is None:
                        # 尝试从标题样式名称中提取级别数字
                        level_match = re.search(r'\d+', style_name)
                        if level_match:
                            level = int(level_match.group())
                            if 1 <= level <= 9:
                                outline_level = level - 1  # 转换为0-8范围
            except (AttributeError, ValueError):
                pass  # 忽略获取样式时可能出现的错误
            
            # 如果段落有大纲级别或是标题样式，则不进行首行缩进
            if outline_level is not None or is_heading_style:
                level = outline_level + 1  # 大纲级别0-8对应标题级别1-9
                self._log(f"段落 {current_block_num}: 大纲级别 {level} 标题 - \"{para_text_preview}...\"")
                self._strip_leading_whitespace(para)
                
                # 根据大纲级别应用不同的字体和格式
                if level == 1:
                    # 检查配置中是否有h1_bold属性，如果没有则默认为False
                    h1_bold = self.config.get('h1_bold', False)
                    self._apply_font_to_runs(para, self.config['h1_font'], self.config['h1_size'], set_color=apply_color, is_bold=h1_bold)
                    # 应用一级标题段前、段后间距（直接使用磅值）
                    h1_space_before_pts = float(self.config['h1_space_before'])
                    h1_space_after_pts = float(self.config['h1_space_after'])
                    para.paragraph_format.space_before = Pt(h1_space_before_pts)
                    para.paragraph_format.space_after = Pt(h1_space_after_pts)
                elif level == 2:
                    # 检查配置中是否有h2_bold属性，如果没有则默认为True
                    h2_bold = self.config.get('h2_bold', True)
                    self._apply_font_to_runs(para, self.config['h2_font'], self.config['h2_size'], set_color=apply_color, is_bold=h2_bold)
                    # 应用二级标题段前、段后间距（直接使用磅值）
                    h2_space_before_pts = float(self.config['h2_space_before'])
                    h2_space_after_pts = float(self.config['h2_space_after'])
                    para.paragraph_format.space_before = Pt(h2_space_before_pts)
                    para.paragraph_format.space_after = Pt(h2_space_after_pts)
                elif level == 3:
                    # 检查配置中是否有h3_bold属性，如果没有则默认为False
                    h3_bold = self.config.get('h3_bold', False)
                    # 三级标题样式和间距设置
                    self._apply_font_to_runs(para, self.config['h3_font'], self.config['h3_size'], set_color=apply_color, is_bold=h3_bold)
                    # 应用三级标题段前、段后间距（直接使用磅值）
                    h3_space_before_pts = float(self.config['h3_space_before'])
                    h3_space_after_pts = float(self.config['h3_space_after'])
                    para.paragraph_format.space_before = Pt(h3_space_before_pts)
                    para.paragraph_format.space_after = Pt(h3_space_after_pts)
                else:
                    # 4-9级标题使用正文字体和默认间距
                    self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                
                # 设置标题行间距
                spacing = para._p.get_or_add_pPr().get_or_add_spacing()
                spacing.set(qn('w:beforeAutospacing'), '0')
                spacing.set(qn('w:afterAutospacing'), '0')
                para.paragraph_format.line_spacing = Pt(self.config['line_spacing'])
                
                # 标题不缩进
                self._apply_text_indent_and_align(para)
                self._reset_pagination_properties(para)
                block_idx += 1
                continue

            # 取消自动识别"一、"、"（一）"、"1."、"(1)"等四级常规标题的功能
            # 直接将这些段落作为正文处理
            if re_h1.match(text_to_check) or re_h2.match(text_to_check) or re_h3.match(text_to_check) or re_h4.match(text_to_check):
                self._log(f"段落 {current_block_num}: 常规标题格式文本 - \"{para_text_preview}...\" (已禁用自动识别，按正文处理)")
            
            # 所有段落都按正文处理
            self._log(f"段落 {current_block_num}: 正文 - \"{para_text_preview}...\"")
            self._strip_leading_whitespace(para)
            self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)
            # 正文需要首行缩进
            self._apply_body_text_indent_and_align(para)
            self._reset_pagination_properties(para)
            
            block_idx += 1
        
        self._apply_page_setup(doc, is_from_txt=is_from_txt)
        self._log("正在保存最终文档...")
        doc.save(output_path)


class WordFormatterGUI:
    def __init__(self, master):
        self.master = master
        master.title("报告自动排版工具_JXSLY V1.0.0")
        # 增加窗体尺寸：宽度增加7%，高度再增加5%
        # 原始尺寸：1320x813，调整后约为1412x942
        master.geometry("1412x942")
        master.minsize(1200, 700)  # 设置最小窗口大小
        
        # 使程序启动时界面位于屏幕中央
        # 先更新窗口任务，确保窗口尺寸已应用
        master.update_idletasks()
        # 获取屏幕尺寸
        screen_width = master.winfo_screenwidth()
        screen_height = master.winfo_screenheight()
        # 获取窗口尺寸
        window_width = 1412
        window_height = 942
        # 计算居中位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        # 设置窗口位置
        master.geometry(f'{window_width}x{window_height}+{x}+{y}')

        self.font_size_map = {
            '一号 (26pt)': 26, '小一 (24pt)': 24, '二号 (22pt)': 22, '小二 (18pt)': 18,
            '三号 (16pt)': 16, '小三 (15pt)': 15, '四号 (14pt)': 14, '小四 (12pt)': 12,
            '五号 (10.5pt)': 10.5, '小五 (9pt)': 9
        }
        self.font_size_map_rev = {v: k for k, v in self.font_size_map.items()}
        
        self.default_params = {
            'page_number_align': '奇偶分页', 'line_spacing': 28,
            'margin_top': 3.7, 'margin_bottom': 3.5, 
            'margin_left': 2.8, 'margin_right': 2.6,
            'h1_font': '黑体', 'h2_font': '楷体_GB2312', 'h3_font': '宋体', 'body_font': '仿宋_GB2312',
            'page_number_font': '宋体', 'table_caption_font': '黑体', 'figure_caption_font': '黑体',
            'h1_size': 16, 'h1_space_before': 24, 'h1_space_after': 24,
            'h2_size': 16, 'h2_space_before': 24, 'h2_space_after': 24,
            'h3_size': 12, 'h3_space_before': 24, 'h3_space_after': 24,
            'body_size': 16, 'page_number_size': 14,
            'table_caption_size': 14, 'figure_caption_size': 14,
            # 添加表格标题和图表标题的大纲级别设置，默认为6级
            'table_caption_outline_level': 8, 'figure_caption_outline_level': 6,
            'set_outline': True,
            # 添加标题粗体设置
            'h1_bold': False,  # 一级标题默认不加粗
            'h2_bold': True,   # 二级标题默认加粗
            'h3_bold': False,  # 三级标题默认不加粗
            'table_caption_bold': False,  # 表格标题默认不加粗
            'figure_caption_bold': False  # 图形标题默认不加粗
        }
        self.font_options = {
            'h1': ['黑体', '方正黑体_GBK', '方正黑体简体', '华文黑体', '宋体'],
            'h2': ['楷体_GB2312', '方正楷体_GBK', '楷体', '方正楷体简体', '华文楷体', '宋体'],
            'h3': ['宋体', '仿宋_GB2312', '方正仿宋_GBK', '仿宋', '方正仿宋简体', '华文仿宋'],
            'body': ['仿宋_GB2312', '方正仿宋_GBK', '仿宋', '方正仿宋简体', '华文仿宋', '宋体'], 
            'table_caption': ['黑体', '宋体', '仿宋_GB2312', '仿宋'], 'figure_caption': ['黑体', '宋体', '仿宋_GB2312', '仿宋']
        }
        self.set_outline_var = tk.BooleanVar(value=self.default_params['set_outline'])

        self.entries = {}
        self.checkboxes = {}  # 存储复选框变量
        
        self.default_config_path = "default_config.json"
        
        self.create_menu()
        self.create_widgets()
        self.load_initial_config()

        self.master.after(250, self.set_initial_pane_position)

    def set_initial_pane_position(self):
        # 获取窗口总宽度，设置左侧占约30%
        total_width = self.master.winfo_width()
        
        if total_width > 100:  # 确保窗口已经渲染
            left_width = int(total_width * 0.3)  # 左侧占30%
            # 使用保存的main_pane引用直接设置位置
            try:
                if hasattr(self, 'main_pane'):
                    self.main_pane.sashpos(0, left_width)
            except Exception as e:
                # 如果直接设置失败，回退到原方法
                for widget in self.master.winfo_children():
                    if isinstance(widget, ttk.PanedWindow):
                        widget.sashpos(0, left_width)
                        break



    def create_menu(self):


        menubar = Menu(self.master)
        # 删除帮助菜单
        self.master.config(menu=menubar)

    def create_widgets(self):
        # 创建主容器，使用垂直布局
        content_frame = ttk.Frame(self.master)
        content_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 创建水平分割的主面板（上方部分）
        main_pane = ttk.PanedWindow(content_frame, orient=tk.HORIZONTAL)
        main_pane.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        # 保存main_pane引用，便于后续访问
        self.main_pane = main_pane

        # 左侧文件处理区域
        left_frame = ttk.Frame(main_pane)
        main_pane.add(left_frame, weight=3)

        notebook = ttk.Notebook(left_frame)
        notebook.pack(fill=tk.BOTH, expand=True)
        self.notebook = notebook

        file_tab = ttk.Frame(notebook)
        notebook.add(file_tab, text=' 文件批量处理 ')
        
        # 创建统一的内容区域，优化布局减少空白
        left_content_frame = ttk.Frame(file_tab)
        left_content_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 文件列表区域
        list_frame = ttk.LabelFrame(left_content_frame, text="待处理文件列表（可拖拽文件或文件夹）", padding=5)
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        # 文件列表和滚动条
        list_inner_frame = ttk.Frame(list_frame)
        list_inner_frame.pack(fill=tk.BOTH, expand=True)
        
        scrollbar = ttk.Scrollbar(list_inner_frame, orient=tk.VERTICAL)
        # 为文件列表设置固定高度，避免占用过多空间
        self.file_listbox = tk.Listbox(list_inner_frame, yscrollcommand=scrollbar.set, selectmode=tk.EXTENDED)
        scrollbar.config(command=self.file_listbox.yview)
        self.file_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, pady=(0, 5))
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y, pady=(0, 5))
        
        self.file_listbox.drop_target_register(DND_FILES)
        self.file_listbox.dnd_bind('<<Drop>>', self.handle_drop)
        self.placeholder_label = ttk.Label(self.file_listbox, text="可以拖拽文件或文件夹到这里", foreground="grey")
        
        # 文件操作按钮区域
        file_button_frame = ttk.Frame(left_content_frame)
        file_button_frame.pack(fill=tk.X, pady=(5, 0))
        
        # 使用网格布局优化按钮排列
        ttk.Button(file_button_frame, text="添加文件", command=self.add_files).grid(row=0, column=0, sticky='ew', padx=2, pady=2)
        ttk.Button(file_button_frame, text="添加文件夹", command=self.add_folder).grid(row=0, column=1, sticky='ew', padx=2, pady=2)
        ttk.Button(file_button_frame, text="移除文件", command=self.remove_files).grid(row=1, column=0, sticky='ew', padx=2, pady=2)
        ttk.Button(file_button_frame, text="清空列表", command=self.clear_list).grid(row=1, column=1, sticky='ew', padx=2, pady=2)
        
        file_button_frame.columnconfigure(0, weight=1)
        file_button_frame.columnconfigure(1, weight=1)

        # 右侧参数设置区域
        right_frame = ttk.Frame(main_pane, padding=(5, 0, 0, 0))
        main_pane.add(right_frame, weight=7)
        
        # 在主面板下方创建处理日志区域
        log_frame = ttk.LabelFrame(content_frame, text="处理日志", padding=5)
        log_frame.pack(fill=tk.BOTH, expand=False)
        # 确保调试日志文本框能完全拉伸至窗体边缘
        # 限制调试日志面板高度，仅显示必要内容
        self.debug_text = scrolledtext.ScrolledText(log_frame, height=4, state='disabled', wrap=tk.WORD)
        self.debug_text.pack(fill=tk.BOTH, expand=True)
        
        # 创建统一的右侧内容区域，与左侧面板结构保持一致
        right_content_frame = ttk.Frame(right_frame)
        right_content_frame.pack(fill=tk.BOTH, expand=True)
        
        # 创建带滚动条的参数设置区域
        canvas = tk.Canvas(right_content_frame)
        v_scrollbar = ttk.Scrollbar(right_content_frame, orient=tk.VERTICAL, command=canvas.yview)
        canvas.configure(yscrollcommand=v_scrollbar.set)
        
        # 创建参数容器
        params_container = ttk.Frame(canvas)
        canvas_window = canvas.create_window((0, 0), window=params_container, anchor='nw', width=right_content_frame.winfo_width()-20)
        
        # 参数设置框架
        params_frame = ttk.LabelFrame(params_container, text="参数设置", padding=10)
        params_frame.pack(fill=tk.BOTH, expand=True)
        params_frame.columnconfigure(1, weight=1)
        params_frame.columnconfigure(3, weight=1)
        params_frame.columnconfigure(5, weight=1)

        # Helper functions for creating widgets
        def create_entry(label, var_name, r, c, width=12):
            ttk.Label(params_frame, text=label).grid(row=r, column=c, sticky=tk.W, padx=5, pady=3)
            entry = ttk.Entry(params_frame, width=width)
            entry.grid(row=r, column=c+1, sticky=tk.EW, padx=5, pady=3)
            self.entries[var_name] = entry
            return entry
        
        def create_combo(label, var_name, opts, r, c, readonly=True, width=15): 
            ttk.Label(params_frame, text=label).grid(row=r, column=c, sticky=tk.W, padx=5, pady=3)
            state = 'readonly' if readonly else 'normal'
            combo = ttk.Combobox(params_frame, values=opts, state=state, width=width)
            combo.grid(row=r, column=c+1, sticky=tk.EW, padx=5, pady=3)
            self.entries[var_name] = combo
            return combo

        def create_font_size_combo(label, var_name, r, c, width=15):
            ttk.Label(params_frame, text=label).grid(row=r, column=c, sticky=tk.W, padx=5, pady=3)
            combo = ttk.Combobox(params_frame, values=list(self.font_size_map.keys()), width=width)
            combo.grid(row=r, column=c+1, sticky=tk.EW, padx=5, pady=3)
            self.entries[var_name] = combo
            return combo
        
        def create_checkbox(label, var_name, r, c, default_value=False):
            ttk.Label(params_frame, text=label).grid(row=r, column=c, sticky=tk.W, padx=5, pady=3)
            checkbox_var = tk.BooleanVar(value=default_value)
            checkbox = ttk.Checkbutton(params_frame, variable=checkbox_var)
            checkbox.grid(row=r, column=c+1, sticky=tk.W, padx=5, pady=3)
            self.checkboxes[var_name] = checkbox_var
            return checkbox_var
        
        def create_section_header(text, help_text, r):
            header_frame = ttk.Frame(params_frame)
            header_frame.grid(row=r, column=0, columnspan=6, sticky='ew', pady=(15, 5))
            ttk.Label(header_frame, text=text, font=('Helvetica', 10, 'bold')).pack(side=tk.LEFT)
            # 删除帮助提示功能
            ttk.Separator(params_frame, orient='horizontal').grid(row=r+1, column=0, columnspan=6, sticky='ew', pady=(5, 10))
            return r + 2

        row = 0
        
        # Section: Page Layout
        row = create_section_header("页面设置", None, row)
        create_entry("上边距(cm)", 'margin_top', row, 0, width=15)
        create_entry("下边距(cm)", 'margin_bottom', row, 2, width=15)
        row += 1
        create_entry("左边距(cm)", 'margin_left', row, 0, width=15)
        create_entry("右边距(cm)", 'margin_right', row, 2, width=15)
        row += 1

        # Section: Document Title

        # Section: Body and Headings
        row = create_section_header("正文与层级", None, row)
        create_combo("一级标题字体", 'h1_font', self.font_options['h1'], row, 0, readonly=False, width=18)
        create_font_size_combo("一级标题字号", 'h1_size', row, 2, width=18)
        create_checkbox("一级标题加粗", 'h1_bold', row, 4, default_value=False)  # 一级标题默认不加粗
        row += 1
        create_entry("一级段前(磅)", 'h1_space_before', row, 0, width=15)
        create_entry("一级段后(磅)", 'h1_space_after', row, 2, width=15)
        row += 1
        create_combo("二级标题字体", 'h2_font', self.font_options['h2'], row, 0, readonly=False, width=18)
        create_font_size_combo("二级标题字号", 'h2_size', row, 2, width=18)
        create_checkbox("二级标题加粗", 'h2_bold', row, 4, default_value=True)  # 二级标题默认加粗
        row += 1
        create_entry("二级段前(磅)", 'h2_space_before', row, 0, width=15)
        create_entry("二级段后(磅)", 'h2_space_after', row, 2, width=15)
        row += 1
        create_combo("三级标题字体", 'h3_font', self.font_options['h3'], row, 0, readonly=False, width=18)
        create_font_size_combo("三级标题字号", 'h3_size', row, 2, width=18)
        create_checkbox("三级标题加粗", 'h3_bold', row, 4, default_value=False)  # 三级标题默认不加粗
        row += 1
        create_entry("三级段前(磅)", 'h3_space_before', row, 0, width=15)
        create_entry("三级段后(磅)", 'h3_space_after', row, 2, width=15)
        row += 1
        create_combo("正文/四级字体", 'body_font', self.font_options['body'], row, 0, readonly=False, width=18)
        create_font_size_combo("正文/四级字号", 'body_size', row, 2, width=18)
        create_entry("正文行距(磅)", 'line_spacing', row, 4, width=15)
        row += 1
        
        # Section: Other Elements
        row = create_section_header("其他元素", None, row)
        create_combo("表格标题字体", 'table_caption_font', self.font_options['table_caption'], row, 0, readonly=False, width=18)
        create_font_size_combo("表格标题字号", 'table_caption_size', row, 2, width=18)
        # 添加表格标题大纲级别（移到同一行）
        ttk.Label(params_frame, text="表格标题大纲级别").grid(row=row, column=4, sticky=tk.W, padx=5, pady=3)
        table_outline_combo = ttk.Combobox(params_frame, values=['无', '1', '2', '3', '4', '5', '6', '7', '8', '9'], width=18)
        table_outline_combo.grid(row=row, column=5, sticky=tk.EW, padx=5, pady=3)
        table_outline_combo.set('8')  # 默认为8级
        self.entries['table_caption_outline_level'] = table_outline_combo
        row += 1
        # 添加表格标题加粗复选框（放在大纲级别控件下方）
        ttk.Label(params_frame, text="表格标题加粗").grid(row=row, column=4, sticky=tk.W, padx=5, pady=3)
        table_bold_var = tk.BooleanVar(value=False)  # 默认为不加粗
        table_bold_checkbox = ttk.Checkbutton(params_frame, variable=table_bold_var)
        table_bold_checkbox.grid(row=row, column=5, sticky=tk.W, padx=5, pady=3)
        self.checkboxes['table_caption_bold'] = table_bold_var
        row += 1
        
        create_combo("图形标题字体", 'figure_caption_font', self.font_options['figure_caption'], row, 0, readonly=False, width=18)
        create_font_size_combo("图形标题字号", 'figure_caption_size', row, 2, width=18)
        # 添加图表标题大纲级别（移到同一行）
        ttk.Label(params_frame, text="图表标题大纲级别").grid(row=row, column=4, sticky=tk.W, padx=5, pady=3)
        figure_outline_combo = ttk.Combobox(params_frame, values=['无', '1', '2', '3', '4', '5', '6', '7', '8', '9'], width=18)
        figure_outline_combo.grid(row=row, column=5, sticky=tk.EW, padx=5, pady=3)
        figure_outline_combo.set('6')  # 默认为6级
        self.entries['figure_caption_outline_level'] = figure_outline_combo
        row += 1
        # 添加图形标题加粗复选框（放在大纲级别控件下方）
        ttk.Label(params_frame, text="图形标题加粗").grid(row=row, column=4, sticky=tk.W, padx=5, pady=3)
        figure_bold_var = tk.BooleanVar(value=False)  # 默认为不加粗
        figure_bold_checkbox = ttk.Checkbutton(params_frame, variable=figure_bold_var)
        figure_bold_checkbox.grid(row=row, column=5, sticky=tk.W, padx=5, pady=3)
        self.checkboxes['figure_caption_bold'] = figure_bold_var
        row += 1

        # Section: Global Options
        ttk.Separator(params_frame, orient='horizontal').grid(row=row, column=0, columnspan=6, sticky='ew', pady=10)
        row += 1

        # 按钮区域
        button_frame = ttk.Frame(params_container, padding=(0, 10, 0, 10))
        button_frame.pack(fill=tk.X)
        
        # 配置按钮 - 2x2布局
        config_buttons = ttk.LabelFrame(button_frame, text="参数管理", padding=10)
        config_buttons.pack(fill=tk.X, pady=(0, 10))
        ttk.Button(config_buttons, text="加载参数", command=self.load_config).grid(row=0, column=0, sticky='ew', padx=5, pady=5)
        ttk.Button(config_buttons, text="保存参数", command=self.save_config).grid(row=0, column=1, sticky='ew', padx=5, pady=5)
        ttk.Button(config_buttons, text="保存为默认", command=self.save_default_config).grid(row=1, column=0, sticky='ew', padx=5, pady=5)
        ttk.Button(config_buttons, text="恢复内置默认", command=self.load_defaults).grid(row=1, column=1, sticky='ew', padx=5, pady=5)
        config_buttons.columnconfigure(0, weight=1)
        config_buttons.columnconfigure(1, weight=1)

        # 开始排版按钮
        style = ttk.Style()
        style.configure('Success.TButton', font=('Helvetica', 11, 'bold'))
        start_button_frame = ttk.Frame(button_frame)
        # 向下移动1cm（约38像素）
        start_button_frame.pack(fill=tk.X, pady=(38, 0))
        ttk.Button(start_button_frame, text="开始排版", style='Success.TButton', command=self.start_processing).pack(fill=tk.X, ipady=10)

        # 配置Canvas滚动
        def on_canvas_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
            # 调整Canvas内容宽度以适应Canvas
            canvas_width = event.width
            canvas.itemconfig(canvas_window, width=canvas_width-20)

        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))

        canvas.bind('<Configure>', on_canvas_configure)
        params_container.bind('<Configure>', on_frame_configure)
        
        # 添加鼠标滚轮支持
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        canvas.bind_all("<MouseWheel>", on_mousewheel)
        params_container.bind_all("<MouseWheel>", on_mousewheel)
        
        # 布局Canvas和滚动条
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self._update_listbox_placeholder()
        
        # 添加定时器，延迟一小段时间后再次应用默认配置，确保UI控件完全创建
        self.master.after(100, self._apply_default_spacing_values)
    
    def log_to_debug_window(self, message):
        self.master.update_idletasks()
        self.debug_text.config(state='normal')
        self.debug_text.insert(tk.END, message + '\n')
        self.debug_text.config(state='disabled')
        self.debug_text.see(tk.END)
    
    def _apply_default_spacing_values(self):
        # 直接设置标题字体和字号
        if 'h3_font' in self.entries:
            self.entries['h3_font'].set(self.default_params['h3_font'])
        if 'h3_size' in self.entries:
            display_val = self.font_size_map_rev.get(self.default_params['h3_size'], str(self.default_params['h3_size']))
            self.entries['h3_size'].set(display_val)
        
        # 直接设置标题间距输入框的值
        if 'h1_space_before' in self.entries:
            self.entries['h1_space_before'].delete(0, tk.END)
            self.entries['h1_space_before'].insert(0, str(self.default_params['h1_space_before']))
        if 'h1_space_after' in self.entries:
            self.entries['h1_space_after'].delete(0, tk.END)
            self.entries['h1_space_after'].insert(0, str(self.default_params['h1_space_after']))
        if 'h2_space_before' in self.entries:
            self.entries['h2_space_before'].delete(0, tk.END)
            self.entries['h2_space_before'].insert(0, str(self.default_params['h2_space_before']))
        if 'h2_space_after' in self.entries:
            self.entries['h2_space_after'].delete(0, tk.END)
            self.entries['h2_space_after'].insert(0, str(self.default_params['h2_space_after']))
        if 'h3_space_before' in self.entries:
            self.entries['h3_space_before'].delete(0, tk.END)
            self.entries['h3_space_before'].insert(0, str(self.default_params['h3_space_before']))
        if 'h3_space_after' in self.entries:
            self.entries['h3_space_after'].delete(0, tk.END)
            self.entries['h3_space_after'].insert(0, str(self.default_params['h3_space_after']))
        
        # 确认已设置的值 - 不再输出到日志窗口
        # self.log_to_debug_window("标题间距值已设置到输入框:")
        # for key in ['h1_space_before', 'h1_space_after', 'h2_space_before', 'h2_space_after', 'h3_space_before', 'h3_space_after']:
        #     if key in self.entries:
        #         self.log_to_debug_window(f"{key}: {self.entries[key].get()}")
        # 直接设置标题字体和字号
        if 'h3_font' in self.entries:
            self.entries['h3_font'].set(self.default_params['h3_font'])
        if 'h3_size' in self.entries:
            display_val = self.font_size_map_rev.get(self.default_params['h3_size'], str(self.default_params['h3_size']))
            self.entries['h3_size'].set(display_val)
        
        # 直接设置标题间距输入框的值
        if 'h1_space_before' in self.entries:
            self.entries['h1_space_before'].delete(0, tk.END)
            self.entries['h1_space_before'].insert(0, str(self.default_params['h1_space_before']))
        if 'h1_space_after' in self.entries:
            self.entries['h1_space_after'].delete(0, tk.END)
            self.entries['h1_space_after'].insert(0, str(self.default_params['h1_space_after']))
        if 'h2_space_before' in self.entries:
            self.entries['h2_space_before'].delete(0, tk.END)
            self.entries['h2_space_before'].insert(0, str(self.default_params['h2_space_before']))
        if 'h2_space_after' in self.entries:
            self.entries['h2_space_after'].delete(0, tk.END)
            self.entries['h2_space_after'].insert(0, str(self.default_params['h2_space_after']))
        if 'h3_space_before' in self.entries:
            self.entries['h3_space_before'].delete(0, tk.END)
            self.entries['h3_space_before'].insert(0, str(self.default_params['h3_space_before']))
        if 'h3_space_after' in self.entries:
            self.entries['h3_space_after'].delete(0, tk.END)
            self.entries['h3_space_after'].insert(0, str(self.default_params['h3_space_after']))
        
        # 确认已设置的值 - 不再输出到日志窗口
        # self.log_to_debug_window("标题间距值已设置到输入框:")
        # for key in ['h1_space_before', 'h1_space_after', 'h2_space_before', 'h2_space_after', 'h3_space_before', 'h3_space_after']:
        #     if key in self.entries:
        #         self.log_to_debug_window(f"{key}: {self.entries[key].get()}")

    def load_initial_config(self):
        if os.path.exists(self.default_config_path):
            try:
                with open(self.default_config_path, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                self._apply_config(config)
                # 不再输出默认配置文件加载信息
            except Exception as e:
                self.log_to_debug_window(f"加载默认配置 '{self.default_config_path}' 失败: {e}。将使用内置默认值。")
                self.load_defaults()
        else:
            self.log_to_debug_window("未找到默认配置文件，将使用内置默认值。")
            self.load_defaults()
        
        # 添加定时器，延迟一小段时间后再次应用默认配置，确保UI控件完全创建
        self.master.after(100, self._apply_default_spacing_values)
    
    def _apply_config(self, loaded_config):
        self.set_outline_var.set(loaded_config.get('set_outline', True))
        for key, value in loaded_config.items():
            if key in ['set_outline']: continue
            
            # 处理输入框和下拉框的值
            widget = self.entries.get(key)
            if widget:
                if "_size" in key:
                    display_val = self.font_size_map_rev.get(value, str(value))
                    widget.set(display_val)
                elif isinstance(widget, ttk.Combobox):
                    widget.set(value)
                else:
                    widget.delete(0, tk.END)
                    widget.insert(0, str(value))
            
            # 处理复选框的值（标题粗体设置）
            checkbox_var = self.checkboxes.get(key)
            if checkbox_var is not None:
                checkbox_var.set(bool(value))

    def load_defaults(self):
        self._apply_config(self.default_params)
    
    def collect_config(self):
        config = {}
        # 收集输入框和下拉框的值
        for key, widget in self.entries.items():
            value = widget.get().strip()
            if "_size" in key:
                if value in self.font_size_map:
                    config[key] = self.font_size_map[value]
                else:
                    try: config[key] = float(value)
                    except (ValueError, TypeError):
                        self.log_to_debug_window(f"警告: 无效的字号值 '{value}' for '{key}'. 使用默认值 16pt。")
                        config[key] = 16
            else:
                try: config[key] = float(value) if '.' in value else int(value)
                except (ValueError, TypeError): config[key] = value
        # 收集复选框的值（标题粗体设置）
        for key, checkbox_var in self.checkboxes.items():
            config[key] = checkbox_var.get()
        config['set_outline'] = self.set_outline_var.get()
        return config

    def save_config(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("JSON files", "*.json")])
        if file_path:
            with open(file_path, 'w', encoding='utf-8') as f: json.dump(self.collect_config(), f, ensure_ascii=False, indent=4)
            messagebox.showinfo("成功", f"配置已保存至 {file_path}")
    
    def save_default_config(self):
        try:
            with open(self.default_config_path, 'w', encoding='utf-8') as f:
                json.dump(self.collect_config(), f, ensure_ascii=False, indent=4)
            messagebox.showinfo("成功", f"当前配置已保存为默认配置。\n下次启动软件时将自动加载。")
        except Exception as e:
            messagebox.showerror("错误", f"保存默认配置失败: {e}")

    def load_config(self):
        file_path = filedialog.askopenfilename(filetypes=[("JSON files", "*.json")])
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    loaded_config = json.load(f)
                self._apply_config(loaded_config)
                messagebox.showinfo("成功", "配置已加载")
            except Exception as e:
                messagebox.showerror("错误", f"加载参数文件失败: {e}")

    def _update_listbox_placeholder(self):
        if self.file_listbox.size() == 0:
            self.placeholder_label.place(in_=self.file_listbox, relx=0.5, rely=0.5, anchor=tk.CENTER)
        else:
            self.placeholder_label.place_forget()

    def handle_drop(self, event):
        paths = self.master.tk.splitlist(event.data)
        self._add_paths_to_listbox(paths)

    def _add_paths_to_listbox(self, paths):
        current_files = set(self.file_listbox.get(0, tk.END))
        added_count = 0
        
        for path in paths:
            if os.path.isdir(path):
                for root, _, files in os.walk(path):
                    for f in files:
                        if f.lower().endswith(('.docx', '.doc', '.wps', '.txt')):
                            full_path = os.path.join(root, f)
                            if full_path not in current_files:
                                self.file_listbox.insert(tk.END, full_path)
                                current_files.add(full_path)
                                added_count += 1
            elif os.path.isfile(path):
                if path.lower().endswith(('.docx', '.doc', '.wps', '.txt')):
                    if path not in current_files:
                        self.file_listbox.insert(tk.END, path)
                        current_files.add(path)
                        added_count += 1
        
        if added_count > 0:
            self.log_to_debug_window(f"通过按钮或拖拽添加了 {added_count} 个新文件。")
        
        self._update_listbox_placeholder()

    def add_files(self):
        files = filedialog.askopenfilenames(filetypes=[("所有支持的文件", "*.docx;*.doc;*.wps;*.txt"), ("Word 文档", "*.docx;*.doc"), ("WPS 文档", "*.wps"), ("纯文本", "*.txt")])
        if files:
            self._add_paths_to_listbox(files)
        
    def add_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            self._add_paths_to_listbox([folder])

    def remove_files(self):
        selected_indices = self.file_listbox.curselection()
        if not selected_indices:
            messagebox.showinfo("提示", "请先在列表中选择要移除的文件。")
            return
        for index in sorted(selected_indices, reverse=True):
            self.file_listbox.delete(index)
        self._update_listbox_placeholder()

    def clear_list(self): 
        self.file_listbox.delete(0, tk.END)
        self._update_listbox_placeholder()

    def start_processing(self):
        warning_title = "处理前重要提示"
        warning_message = (
            "为了防止数据丢失，请在继续前关闭所有已打开的Word和WPS文档（包括wps、表格、PPT等所有文档）。\n\n"
            "本程序在转换文件格式时需要调用Word/WPS程序，这可能会导致您未保存的工作被强制关闭。\n\n"
            "您确定要继续吗？"
        )
        if not messagebox.askokcancel(warning_title, warning_message):
            self.log_to_debug_window("用户已取消操作。")
            return
            
        self.debug_text.config(state='normal'); self.debug_text.delete('1.0', tk.END); self.debug_text.config(state='disabled')
        
        processor = WordProcessor(self.collect_config(), self.log_to_debug_window)
        active_tab_index = self.notebook.index(self.notebook.select())

        try:
            if active_tab_index == 0:
                file_list = self.file_listbox.get(0, tk.END)
                if not file_list:
                    messagebox.showwarning("警告", "文件列表为空，请先添加文件！"); return
                output_dir = filedialog.askdirectory(title="请选择一个文件夹用于存放处理后的文件")
                if not output_dir: return

                success_count, fail_count = 0, 0
                for i, input_path in enumerate(file_list):
                    try:
                        self.log_to_debug_window(f"\n--- 开始处理文件 {i+1}/{len(file_list)}: {os.path.basename(input_path)} ---")
                        base_name = os.path.splitext(os.path.basename(input_path))[0]
                        output_path = os.path.join(output_dir, f"{base_name}_formatted.docx")
                        processor.format_document(input_path, output_path)
                        self.log_to_debug_window(f"✅ 文件处理成功，已保存至: {output_path}")
                        success_count += 1
                    except Exception as e:
                        logging.error(f"处理文件失败: {input_path}\n{e}", exc_info=True)
                        self.log_to_debug_window(f"\n❌ 处理文件 {os.path.basename(input_path)} 时发生严重错误：\n{e}")
                        fail_count += 1
                    finally:
                        processor._cleanup_temp_files()
                
                summary_message = f"批量处理完成！\n\n成功: {success_count}个\n失败: {fail_count}个"
                if fail_count > 0: summary_message += "\n\n失败详情请查看日志窗口。"
                messagebox.showinfo("完成", summary_message)
                self.log_to_debug_window(f"\n🎉 {summary_message}")
                self.log_to_debug_window("\n💡 提示：处理完成的文件可能正在被系统占用，请稍等几秒后再打开。")

        
        except Exception as e:
            logging.error(f"处理过程中发生严重错误: {e}", exc_info=True)
            self.log_to_debug_window(f"\n❌ 处理过程中发生严重错误：\n{e}")
            messagebox.showerror("错误", f"处理过程中发生错误：\n{e}")
        finally:
            processor.quit_com_app()
            self.log_to_debug_window("\n💡 所有任务完成，WPS/Word应用已关闭，现在可以安全地打开处理后的文件了。")

if __name__ == "__main__":
    root = TkinterDnD.Tk()
    app = WordFormatterGUI(root)
    root.mainloop()

















































